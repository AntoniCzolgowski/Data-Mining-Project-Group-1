#!/usr/bin/env python3
"""Build CLASSIFICATION_EXTRAS.docx — Practical Playbook + Research Q&A + Ensemble (for final report)."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, 'docs', 'images', 'methods')
OUT = os.path.join(BASE, 'CLASSIFICATION_EXTRAS.docx')


def img(name):
    p = os.path.join(FIG, name)
    return p if os.path.exists(p) else None


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def add_figure(doc, filename, caption, width=Inches(6.2)):
    path = img(filename)
    if path is None:
        doc.add_paragraph(f'[Figure not found: {filename}]')
        return
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def bold_para(doc, bold_text, normal_text=''):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


# ═══════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Classification Analysis \u2014 Extra Sections\nfor Final Report Integration')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Sections removed from M3 report for brevity.\nContents: Ensemble Stacking, Practical Playbook, Research Q&A.\n\nAntoni Czolgowski \u2014 March 2026')
run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# ENSEMBLE STACKING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Ensemble Stacking \u2014 Can We Do Better by Combining Models?', level=1)

doc.add_paragraph(
    'Ensemble stacking combines multiple models together. "Soft Voting" averages the predictions '
    'from RF, XGBoost, and SVM. "Stacking" trains a second-level model '
    '(Logistic Regression) to learn the optimal way to combine the first-level predictions.'
)

add_table(doc,
    ['Model', 'Accuracy', 'F1', 'AUC', 'Type'],
    [
        ['Random Forest', '0.780', '0.718', '0.828', 'Individual (best)'],
        ['Soft Voting (RF+XGB+SVM)', '0.768', '0.691', '0.826', 'Ensemble'],
        ['XGBoost', '0.748', '0.667', '0.803', 'Individual'],
        ['SVM (RBF)', '0.760', '0.663', '0.819', 'Individual'],
        ['Stacking (LR meta)', '0.760', '0.647', '0.816', 'Ensemble'],
        ['Logistic Regression', '0.680', '0.604', '0.745', 'Individual'],
    ]
)

add_figure(doc, 'deepdive2_model_scoreboard.png',
           'Final scoreboard. Left: F1 scores. Right: ROC curves. '
           'Random Forest alone outperforms all ensemble combinations.')

doc.add_paragraph(
    'Ensembles do NOT improve performance. Stacking actually hurts '
    '(F1 drops from 0.718 to 0.647). With only 250 counties, the LR meta-learner '
    'oversmooths predictions. The performance bottleneck is dataset size, not algorithm sophistication.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# PRACTICAL PLAYBOOK
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Practical Playbook \u2014 Democratic Campaign Strategy', level=1)

doc.add_paragraph(
    'This section translates analytical findings into actionable guidance '
    'for Democratic campaign resource allocation in Pennsylvania, Michigan, and North Carolina.'
)

doc.add_heading('The Core Strategic Insight', level=2)
doc.add_paragraph(
    'Electoral volatility is predictable from demographics. Campaigns don\'t need to wait for '
    'polling data \u2014 census data alone can identify high-leverage counties years before an election. '
    'Key: look for racially diverse, economically transitioning communities \u2014 especially suburbs '
    'with rising housing costs.'
)

doc.add_heading('How to Identify High-Priority Counties', level=2)

bold_para(doc, 'Step 1 \u2014 Screen by Diversity: ',
          'Check if the county\'s racial composition has no single group above ~60%. '
          'Counties above this diversity threshold are 7\u00d7 more likely to swing.')

bold_para(doc, 'Step 2 \u2014 Check for Economic Transition: ',
          'Among diverse counties, prioritize those where rents are rising, '
          'education levels are shifting, or there is significant immigration.')

bold_para(doc, 'Step 3 \u2014 Weight by Vote Volume and Margin: ',
          'A volatile county with 800,000 votes and a 0.1% margin (Bucks, PA) '
          'is worth more than one with 15,000 votes and a 7% margin (Scotland, NC).')

bold_para(doc, 'Step 4 \u2014 Tailor by State: ',
          'The same demographics mean different things in different states. '
          'Use state-specific models, not national ones.')

doc.add_heading('State-by-State Strategy', level=2)

bold_para(doc, 'PENNSYLVANIA \u2014 The Philadelphia Suburban Collar')
doc.add_paragraph(
    'PA\'s volatility is driven by racial diversity and housing cost in the suburban belt. '
    'The decisive battleground is the Philadelphia collar: Bucks (0.1% margin, 802k votes), '
    'Chester, Montgomery, plus the Lehigh Valley (Northampton, Lehigh). '
    'Monroe County (171k votes, 0.8% margin) is an emerging target.'
)
doc.add_paragraph('Primary targets: Bucks, Monroe, Northampton, Chester, Lehigh', style='List Bullet')
doc.add_paragraph('Warning: Philadelphia itself is stable despite being diverse \u2014 don\'t conflate suburban and urban dynamics', style='List Bullet')

bold_para(doc, 'MICHIGAN \u2014 Two-Track Targeting')
doc.add_paragraph(
    'Track 1: Small "misfit" counties (Leelanau, Marquette, Isabella, Grand Traverse) \u2014 '
    'university/resort towns that serve as early-warning indicators. '
    'Track 2: Large population centers (Macomb, Oakland, Kent, Genesee) deliver actual votes. '
    'MI volatility is education-driven, not diversity-driven.'
)
doc.add_paragraph('Primary targets (volume): Macomb, Oakland, Kent, Genesee', style='List Bullet')
doc.add_paragraph('Primary targets (signal): Leelanau, Marquette, Grand Traverse, Isabella', style='List Bullet')

bold_para(doc, 'NORTH CAROLINA \u2014 The Southern Realignment')
doc.add_paragraph(
    'NC is undergoing a deeper realignment. 49% of NC counties show Steady Red Drift. '
    'Cabarrus County (120k votes, Charlotte suburb) is the standout target. '
    'NC volatility is poverty-driven \u2014 messaging should address economic conditions.'
)
doc.add_paragraph('Primary targets: Cabarrus (Charlotte suburb), Wake adjacent counties', style='List Bullet')
doc.add_paragraph('Diagnostic: Scotland, Nash, Lenoir (small but indicative)', style='List Bullet')

doc.add_heading('What NOT to Do', level=2)
doc.add_paragraph('Don\'t apply a one-size-fits-all model across states', style='List Bullet')
doc.add_paragraph('Don\'t target low-diversity counties \u2014 they are almost never volatile', style='List Bullet')
doc.add_paragraph('Don\'t conflate urban and suburban dynamics \u2014 cities are stable even when diverse', style='List Bullet')
doc.add_paragraph('Don\'t assume 2020 surge is permanent \u2014 47% of counties showed Blue Bounceback', style='List Bullet')
doc.add_paragraph('Don\'t ignore Steady Red Drift counties (50% high vol, ongoing realignment)', style='List Bullet')

doc.add_heading('The "Money List" \u2014 10 Counties That Could Decide 2028', level=2)
add_table(doc,
    ['Rank', 'County', 'State', 'Why It Matters'],
    [
        ['1', 'Bucks', 'PA', '802k votes, 0.1% margin, diverse suburb \u2014 could decide PA'],
        ['2', 'Lackawanna', 'PA', '233k votes, Scranton area, blue misfit trending volatile'],
        ['3', 'Cabarrus', 'NC', '120k votes, Charlotte suburb, fast-growing and diversifying'],
        ['4', 'Leelanau', 'MI', 'Small but highest misfit score \u2014 bellwether for MI trends'],
        ['5', 'Scotland', 'NC', 'Highest volatility score, diverse-poor rural archetype'],
        ['6', 'Marquette', 'MI', 'University town, votes Dem against all demographics'],
        ['7', 'Genesee', 'MI', '223k votes (Flint area), working-class transition zone'],
        ['8', 'Monroe', 'PA', '171k votes, Poconos, NYC migration transforming politics'],
        ['9', 'Isabella', 'MI', 'University county (CMU), young and volatile'],
        ['10', 'Grand Traverse', 'MI', '63k votes, resort/wine country, 1.7% margin'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# RESEARCH Q&A
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Answering All Research Questions', level=1)

doc.add_heading('Results Application Questions', level=2)

bold_para(doc, 'Q: Which counties illustrate a combination of electoral elasticity and impact (vote volume), '
               'maximizing the marginal utility of Democratic resource allocation?')
doc.add_paragraph(
    'Our campaign priority scoring directly answers this. Bucks County, PA is #1: '
    '802,000 votes, 0.1% margin, Highly Volatile. The full top 10 combines volatility, misfit score, '
    'vote volume, and margin closeness. Among high-volume counties (>100k votes), '
    'Bucks, Lackawanna, Cabarrus, Genesee, and Monroe offer the best return on campaign investment.'
)

bold_para(doc, 'Q: Which swing-state counties had the highest electorate volatility between 2016, 2020, and 2024?')
doc.add_paragraph(
    'The most volatile counties include Scotland (NC), Leelanau (MI), Marquette (MI), '
    'Bucks (PA), Cabarrus (NC), and Lackawanna (PA). "Steady Red Drift" counties are the most '
    'volatile as a group (50% high volatility). In terms of vote volume, Bucks PA (802k), '
    'Genesee MI (223k), and Lackawanna PA (233k) are the highest-volume volatile counties.'
)

bold_para(doc, 'Q: What specific demographic and economic features distinguish "High-Volatility" '
               'clusters from "Stable" clusters?')
doc.add_paragraph(
    'High-volatility counties are distinguished by: (1) Higher racial diversity index '
    '(odds ratio 7.15\u00d7 above threshold), (2) Higher housing costs (rent is the only universal '
    'predictor), (3) Higher poverty rates (especially NC), (4) Higher education levels '
    '(especially MI), (5) Lower homeownership and younger median age. '
    'Stable counties are racially homogeneous, moderate housing costs, higher homeownership, older populations. '
    'SHAP archetype analysis identifies 5 distinct types with near-perfect separation.'
)

doc.add_heading('Hypotheses / Predictions', level=2)

bold_para(doc, 'H1: Electoral volatility is inversely correlated with wealth.')
doc.add_paragraph(
    'CONFIRMED. SHAP analysis: r = \u22120.602 (p < 0.0001) between income and high-volatility. '
    'Poverty rate is top-5 across all models. Nuance: income alone is insufficient '
    '(F1 = 0.570); it interacts with diversity \u2014 wealthy diverse suburbs are volatile, '
    'wealthy homogeneous exurbs are not.'
)

bold_para(doc, 'H2: A composite Racial Diversity Index will demonstrate higher feature importance '
               'than any single demographic.')
doc.add_paragraph(
    'CONFIRMED across 6 evidence streams: (1) #1 permutation importance, '
    '(2) higher SHAP than pct_black (0.397 vs 0.347), (3) largest PDP range (0.14), '
    '(4) race group alone = 86% of baseline, (5) 6/10 top SHAP interactions, '
    '(6) threshold odds ratio = 7.15\u00d7.'
)

bold_para(doc, 'H3: High-volatility counties function as macro-trend amplifiers.')
doc.add_paragraph(
    'PARTIALLY CONFIRMED. "Blue Bounceback" (47%) amplified both the 2020 blue wave and 2024 correction. '
    '"Steady Red Drift" (38%) amplifies the Republican trend. '
    'However, cross-state analysis shows mechanisms are state-specific, '
    'and misfit counties sometimes swing against expectations. '
    'The amplifier hypothesis holds for the majority but is not universal.'
)

doc.add_heading('Further Contextual Analysis', level=2)

bold_para(doc, 'Q: Do demographic predictors generalize across state lines?')
doc.add_paragraph(
    'NO. LOSO F1 = 0.195, indistinguishable from random (0.200). '
    'PA volatility is diversity-driven, MI education-driven, NC poverty-driven. '
    'Only housing cost predicts in all three states. SHAP archetypes confirm: '
    'Cluster 0 (NC-only) and Cluster 2 (MI-heavy) represent non-transferable mechanisms.'
)

bold_para(doc, 'Q: Does electoral volatility exhibit temporal consistency?')
doc.add_paragraph(
    'The swing map shifts significantly. 47% showed Blue Bounceback (blue 2020 \u2192 red 2024). '
    'Only 14% drifted consistently in one direction. Counties with growing Hispanic populations '
    'and rising rents are becoming more volatile over time, '
    'suggesting the swing map evolves dynamically with demographic change.'
)

bold_para(doc, 'Q: Correlation between voter turnout variance and vote share volatility?')
doc.add_paragraph(
    'Our analysis focused on margin volatility rather than turnout variance. '
    'Indirect evidence: log_total_population has moderate importance but works through interactions '
    '(PDP is mostly flat). Both large and small counties can be volatile. '
    'Direct turnout-variance analysis would require year-over-year turnout data \u2014 '
    'potential future extension.'
)

# ── Save ──
doc.save(OUT)
print(f'Saved: {OUT}')