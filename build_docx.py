#!/usr/bin/env python3
"""Build MILESTONE_3_CLASSIFICATION.docx — comprehensive classification report with all visualizations."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ── Paths ──
BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, 'docs', 'images', 'methods')
OUT = os.path.join(BASE, 'MILESTONE_3_CLASSIFICATION.docx')


def img(name):
    """Return full path to a figure, or None if missing."""
    p = os.path.join(FIG, name)
    return p if os.path.exists(p) else None


# ── Document setup ──
doc = Document()

# -- Style tweaks --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption, width=Inches(6.2)):
    """Add a figure with caption."""
    path = img(filename)
    if path is None:
        doc.add_paragraph(f'[Figure not found: {filename}]').italic = True
        return
    doc.add_picture(path, width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def bold_para(doc, bold_text, normal_text=''):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_takeaway_box(doc, text):
    """Add a highlighted takeaway / section conclusion."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('\u25b6 Section Takeaway: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_method_explainer(doc, title, text):
    """Add a plain-language method explanation box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f'\U0001f4d6 What is {title}? ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x2a, 0x6e, 0x2a)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Milestone 3: Classification of Electoral Volatility')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Predicting County-Level Electoral Instability\nfrom Demographic Census Data')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Antoni Czolgowski\n')
run.bold = True
run.font.size = Pt(13)
run = meta.add_run('CSCI 5502 Data Mining \u2014 Spring 2026\nUniversity of Colorado Boulder\n\nMarch 2026')
run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Research Questions & Motivation',
    '2. Dataset & Feature Engineering',
    '3. Data Formatting for Models',
    '4. Model A \u2014 Random Forest',
    '5. Model B \u2014 XGBoost',
    '6. Model C \u2014 SVM (Support Vector Machine)',
    '7. Model D \u2014 Demographic Misfit Detector (Logistic Regression)',
    '8. Binary Classification \u2014 High vs Low Volatility',
    '9. Ensemble Stacking',
    '10. Model Comparison & Performance Evaluation',
    '11. Cross-State Generalization Experiment',
    '12. Hypothesis Testing',
    '13. Advanced Analysis \u2014 Deep-Dives',
    '14. Key Findings & Grand Synthesis',
    '15. Practical Playbook \u2014 Democratic Campaign Strategy',
    '16. Answering All Research Questions',
    '17. Appendix: Figures & Notebooks',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. RESEARCH QUESTIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('1. Research Questions & Motivation', level=1)

doc.add_paragraph(
    'The central question driving this analysis is simple but powerful:'
)
p = doc.add_paragraph()
r = p.add_run(
    'Can we predict which U.S. counties will be electorally volatile \u2014 '
    'that is, likely to swing between parties \u2014 using only demographic data, '
    'without knowing how they voted in the past?'
)
r.bold = True
r.font.size = Pt(12)

doc.add_paragraph(
    'Electoral volatility matters because it determines where campaign resources '
    'have the highest return on investment. A county that always votes 70\u201330 for one party '
    'is not worth contesting. But a county that swings 5\u201310 points between elections '
    'is where the next election will be decided. If we can identify these counties from '
    'their demographics alone, campaigns can target resources before election data confirms the swing.'
)

doc.add_paragraph(
    'We study 250 counties across three presidential battleground states \u2014 '
    'Pennsylvania (67 counties), Michigan (83), and North Carolina (100) \u2014 '
    'using American Community Survey (ACS) census data merged with election results '
    'from 2016, 2020, and 2024. The choice of these three states is deliberate: '
    'they represent different political geographies (Rust Belt, suburban Northeast, '
    'and the changing South) and together decided the 2024 presidential election.'
)

doc.add_heading('Research Hypotheses', level=2)

bold_para(doc, 'H1 \u2014 Wealth-Volatility Inverse: ',
          'We predict that wealthier counties are more electorally stable. '
          'Counties in the bottom quartile of median household income should show '
          'significantly higher variation in partisan swing than counties in the top quartile. '
          'In plain terms: economic security translates to political consistency.')

bold_para(doc, 'H2 \u2014 Diversity Index > Individual Race: ',
          'We predict that a composite racial diversity index (measuring how "mixed" a county is) '
          'will be a stronger predictor of volatility than any single racial demographic like "% Black" '
          'or "% Hispanic." The hypothesis is that it\'s racial mixing \u2014 not the presence of any '
          'particular group \u2014 that makes a county\'s politics less predictable.')

doc.add_heading('Why Classification?', level=2)
doc.add_paragraph(
    'We frame this as a classification problem because campaign strategists don\'t need '
    'a precise volatility number \u2014 they need to know: "Is this county likely to swing significantly?" '
    'This is a categorization task. We implement multiple classifiers from different algorithmic '
    'families (ensemble trees, gradient boosting, kernel methods, linear models) to understand '
    'which approaches best capture the relationship between demographics and electoral instability.'
)

add_takeaway_box(doc,
    'This analysis asks whether demographics alone can predict electoral swings. '
    'We test this across 250 counties in PA, MI, and NC using 4 different machine learning models, '
    'and then go beyond prediction to understand why some counties swing and others don\'t.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. DATASET
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('2. Dataset & Feature Engineering', level=1)

doc.add_heading('2.1 Source Data', level=2)
add_table(doc,
    ['File', 'Rows', 'Description'],
    [
        ['master_dataset.csv', '750 (250 \u00d7 3 years)', '39 demographic + election columns per county-year'],
        ['master_dataset_scaled.csv', '750', 'Contains race_entropy_norm (racial diversity index, 0\u20131)'],
        ['county_volatility_dimTable.csv', '250', 'Volatility target computed from margin changes'],
    ]
)

doc.add_heading('2.2 Target Variable \u2014 What We Are Predicting', level=2)
doc.add_paragraph(
    'Our target variable, vol_z_abs_sum, measures how much a county\'s partisan lean '
    'shifted across two election cycles. In plain terms, it answers: '
    '"How much did this county swing between 2016\u20132020 and 2020\u20132024?"'
)
doc.add_paragraph(
    'We calculate this by taking the margin change in each cycle '
    '(e.g., a county that went from R+10 to R+2 swung 8 points toward Democrats), '
    'converting to z-scores (so we compare counties on a fair scale), '
    'taking absolute values (we care about the size of the swing, not the direction), '
    'and summing across both cycles.'
)
doc.add_paragraph('This was binned into two formats:')
doc.add_paragraph(
    '5-class quintiles (50 counties each): Very Stable, Stable, Moderate, Volatile, Highly Volatile',
    style='List Bullet')
doc.add_paragraph(
    'Binary: High Volatility (Q4 + Q5, n=100) vs Low Volatility (Q1\u2013Q3, n=150)',
    style='List Bullet')

doc.add_heading('2.3 Feature Inventory (28 Demographic Features)', level=2)
doc.add_paragraph(
    'We use 28 demographic features organized into 6 thematic groups. '
    'Critically, all election-derived variables (vote counts, vote shares, margins) '
    'are excluded from the features to prevent data leakage \u2014 we only use census demographics.'
)

add_table(doc,
    ['Group', 'Features', 'Count'],
    [
        ['Race / Ethnicity', 'pct_black, pct_asian, pct_two_or_more_races,\npct_hispanic, pct_non_hispanic_white, race_entropy_norm', '6'],
        ['Urbanization', 'log_total_population, log_population_density,\npct_drive_alone, pct_carpool, pct_public_transit, pct_work_from_home', '6'],
        ['Education', 'pct_hs_or_higher, pct_bachelors_plus', '2'],
        ['Housing', 'log_median_gross_rent, log_median_home_value,\npct_owner_occupied', '3'],
        ['Economic', 'log_median_household_income, pct_below_poverty,\npct_income_under_25k, pct_income_50k_100k, unemployment_rate', '5'],
        ['Household / Age', 'median_age, pct_senior_65plus, pct_young_adult_18_24,\npct_foreign_born, pct_family_households, pct_married_couple, pct_living_alone', '7'],
    ]
)

doc.add_heading('2.4 The race_entropy_norm Feature \u2014 Measuring Racial Mixing', level=2)
doc.add_paragraph(
    'This is a key feature that deserves explanation. race_entropy_norm is a Shannon entropy index '
    'computed from 4 racial categories (Non-Hispanic White, Black, Asian, Other/Two+). '
    'Think of it as a "diversity score":'
)
doc.add_paragraph('A score of 0 means the county is entirely one race (completely homogeneous)', style='List Bullet')
doc.add_paragraph('A score of 1 means the county is perfectly split across all 4 groups', style='List Bullet')
doc.add_paragraph('A score of ~0.55 means no single group exceeds about 60% of the population \u2014 a genuinely mixed community', style='List Bullet')
doc.add_paragraph(
    'The key insight is that this measures mixing, not the presence of any particular group. '
    'A county that is 90% White and 10% Black has low entropy. A county that is 40% White, '
    '30% Black, 20% Hispanic, 10% Asian has high entropy. As we will show, '
    'it is this mixing \u2014 not the share of any group \u2014 that predicts volatility.'
)

doc.add_heading('2.5 Leakage Prevention', level=2)
doc.add_paragraph(
    'To ensure our models genuinely learn from demographics (not from past elections), '
    'we strictly exclude all election variables from features: dem_votes, rep_votes, total_votes, '
    'dem_pct, rep_pct, dem_margin. The volatility target is derived from election data, '
    'but it never appears among the predictor features.'
)

add_takeaway_box(doc,
    'We use 28 purely demographic features (race, income, education, housing, urbanization, age) '
    'to predict electoral volatility across 250 counties. No election data is used as input \u2014 '
    'this is a genuine test of whether demographics alone can explain political instability.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. DATA FORMATTING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('3. Data Formatting for Models', level=1)

doc.add_paragraph(
    'Different machine learning algorithms have different requirements for how data should be '
    'formatted. This section describes the preprocessing steps we took to ensure all models '
    'receive properly formatted input.'
)

doc.add_heading('3.1 Log-Transformation of Skewed Variables', level=2)
doc.add_paragraph(
    'Some variables (like population and income) have extreme outliers \u2014 '
    'Philadelphia has 1.6 million people while some rural counties have 2,000. '
    'These extreme values can distort model training. We applied a logarithmic transformation '
    'to 5 such variables, which compresses the range and brings outliers closer to the bulk of the data.'
)
add_table(doc,
    ['Variable', 'Before (Skewness)', 'After (Skewness)', 'Why Transform?'],
    [
        ['total_population', '5.82', '0.89', 'Few very large counties dominate'],
        ['population_density', '4.61', '0.34', 'Extreme urban outliers (Philadelphia)'],
        ['median_household_income', '1.03', '0.72', 'Income right-tail'],
        ['median_home_value', '1.67', '0.81', 'Housing market outliers'],
        ['median_gross_rent', '0.98', '0.54', 'Rent distribution tail'],
    ]
)

doc.add_heading('3.2 Removing Redundant Features', level=2)
doc.add_paragraph(
    'When two features are highly correlated (r > 0.85), they contain essentially '
    'the same information. Including both wastes model capacity and can cause instability. '
    'We dropped pct_income_over_100k (r = 0.97 with log_median_household_income) since '
    'both capture the same "wealth" signal.'
)

doc.add_heading('3.3 Feature Scaling (StandardScaler)', level=2)
doc.add_paragraph(
    'We standardized all 28 features to have zero mean and unit variance. '
    'This is essential for SVM (which computes distances between data points) and '
    'Logistic Regression (which uses gradient-based optimization). '
    'For Random Forest and XGBoost, scaling is not strictly required, '
    'but we apply it for consistency.'
)

doc.add_heading('3.4 Before-and-After Snapshots', level=2)
add_table(doc,
    ['Stage', 'total_population', 'pct_below_poverty'],
    [
        ['Raw', 'mean=126,437; range [2,144 \u2013 1,603,797]', 'mean=14.8%; range [4.1% \u2013 35.2%]'],
        ['After log-transform', 'mean=10.8; range [7.7 \u2013 14.3]', '(unchanged)'],
        ['After scaling', 'mean=0.0; range [\u22122.6 \u2013 2.9]', 'mean=0.0; range [\u22121.8 \u2013 3.3]'],
    ]
)

doc.add_heading('3.5 Model-Specific Requirements', level=2)
add_table(doc,
    ['Model', 'Needs Numerical?', 'Needs Scaling?', 'Our Approach'],
    [
        ['Random Forest', 'No', 'No', 'Scaling applied for consistency'],
        ['XGBoost', 'No', 'No', 'Labels adjusted to 0-indexed'],
        ['SVM', 'Yes', 'Yes', 'StandardScaler required'],
        ['Logistic Regression', 'Yes', 'Yes', 'StandardScaler + balanced class weights'],
    ]
)

add_takeaway_box(doc,
    'Data was log-transformed (to tame outliers), de-duplicated (removing redundant correlated features), '
    'and scaled (to put all features on equal footing). These steps ensure each model receives '
    'clean, properly formatted input for fair comparison.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. MODEL A — RANDOM FOREST
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('4. Model A \u2014 Random Forest', level=1)

add_method_explainer(doc, 'Random Forest',
    'Imagine asking 300 different analysts to predict county volatility, but each analyst '
    'only sees a random subset of the data and a random subset of the demographic features. '
    'Each analyst builds their own decision tree (a flowchart of yes/no questions like '
    '"Is poverty rate above 15%?"). Then we take a vote across all 300 analysts. '
    'This "wisdom of crowds" approach is Random Forest \u2014 it reduces the chance that any '
    'single quirk in the data misleads the prediction.')

doc.add_heading('Why We Chose It', level=2)
doc.add_paragraph('Does not assume any particular shape of the relationship between demographics and volatility', style='List Bullet')
doc.add_paragraph('Naturally captures interactions (e.g., "diversity matters more in suburban areas")', style='List Bullet')
doc.add_paragraph('Provides built-in feature importance rankings', style='List Bullet')
doc.add_paragraph('Robust to noise \u2014 critical for our relatively small dataset of 250 counties', style='List Bullet')

doc.add_heading('Assumptions', level=2)
doc.add_paragraph('At least some demographic features carry information about volatility', style='List Bullet')
doc.add_paragraph('The 250 counties are reasonably representative of battleground-state counties', style='List Bullet')
doc.add_paragraph('No assumption about the mathematical form of the relationship (linear, curved, etc.)', style='List Bullet')

doc.add_heading('Hyperparameter Tuning', level=2)
doc.add_paragraph(
    'We used a rigorous two-level cross-validation approach to find the best settings. '
    'The outer loop (5 folds) evaluates performance; the inner loop (3 folds) searches for '
    'the best hyperparameters. This prevents the tuning process from inflating our accuracy estimates.'
)
add_table(doc,
    ['Setting', 'Options Tested', 'Best Value', 'What It Controls'],
    [
        ['Number of trees', '100, 200, 300, 500', '300', 'More trees = more stable predictions'],
        ['Maximum depth', '5, 10, 15, 20, unlimited', '8', 'How complex each tree can be'],
        ['Minimum leaf size', '3, 5, 10', '5', 'Prevents trees from memorizing noise'],
        ['Features per split', 'sqrt, log2, 30%, 50%', 'sqrt', 'Diversity among trees'],
        ['Class weighting', 'balanced, subsample, none', 'balanced', 'Handles unequal class sizes'],
    ]
)

doc.add_heading('5-Class Results', level=2)
add_table(doc,
    ['Metric', 'Value', 'What It Means'],
    [
        ['Accuracy', '0.368', '37% of counties correctly classified into exact quintile'],
        ['Precision (macro)', '0.369', 'When it predicts a class, it\'s right 37% of the time'],
        ['Recall (macro)', '0.368', 'It finds 37% of counties in each class'],
        ['F1-macro', '0.362', 'Balanced performance across all 5 classes'],
        ['ROC-AUC', '0.803', 'Good ability to separate high from low volatility'],
    ]
)

doc.add_paragraph(
    'At first glance, 37% accuracy may seem low. But remember: with 5 classes, random guessing '
    'would achieve only 20%. Our model is nearly twice as good as chance. More importantly, '
    'the model excels at identifying the extremes \u2014 Highly Volatile counties (F1=0.59) and '
    'Very Stable counties (F1=0.46) \u2014 which are the most useful for campaign targeting.'
)

add_table(doc,
    ['Very Stable', 'Stable', 'Moderate', 'Volatile', 'Highly Volatile'],
    [['F1 = 0.46', 'F1 = 0.18', 'F1 = 0.28', 'F1 = 0.30', 'F1 = 0.59']]
)

add_figure(doc, 'rf_confusion_matrix.png',
           'Figure 1: Random Forest confusion matrix (5-class). The model is strongest at '
           'identifying the extremes (Q1 and Q5), which are the most actionable for campaigns.')

doc.add_heading('Feature Importance \u2014 What Demographics Matter Most?', level=2)
doc.add_paragraph(
    'Random Forest provides two ways to measure which features matter most:'
)
doc.add_paragraph(
    'Gini importance: How often each feature is used to split data in the decision trees. '
    'This tends to favor features with many possible values.',
    style='List Bullet')
doc.add_paragraph(
    'Permutation importance: How much accuracy drops when we randomly shuffle a feature\'s values. '
    'This is the more reliable measure \u2014 if shuffling a feature destroys accuracy, that feature truly matters.',
    style='List Bullet')

add_figure(doc, 'rf_gini_importance.png',
           'Figure 2: Gini feature importance \u2014 poverty, rent, and diversity index lead the rankings.')

add_figure(doc, 'rf_permutation_importance.png',
           'Figure 3: Permutation importance \u2014 race_entropy_norm (the diversity index) is #1. '
           'Shuffling this feature causes the largest drop in model accuracy.')

doc.add_heading('Challenges & Solutions', level=2)
doc.add_paragraph('Challenge: Only 250 counties (small dataset). Solution: Limited tree depth to 8 and required at least 5 counties per leaf to prevent memorization.', style='List Bullet')
doc.add_paragraph('Challenge: 5 equal-sized classes. Solution: Used balanced class weights to ensure the model pays equal attention to each volatility level.', style='List Bullet')

add_takeaway_box(doc,
    'Random Forest correctly identifies the most volatile and most stable counties '
    'with reasonable accuracy. The single most important demographic feature is the racial '
    'diversity index \u2014 counties with more racial mixing are significantly more likely to swing. '
    'Poverty, housing cost, and household structure also matter.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. MODEL B — XGBOOST
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('5. Model B \u2014 XGBoost', level=1)

add_method_explainer(doc, 'XGBoost (Gradient Boosting)',
    'While Random Forest builds many trees independently and takes a vote, XGBoost builds trees '
    'one at a time, where each new tree specifically tries to fix the mistakes of the previous ones. '
    'Think of it like a student taking a test, getting the wrong answers back, and studying '
    'specifically those topics for the next attempt. This "learning from mistakes" approach '
    'often produces more accurate predictions, especially for hard-to-classify cases.')

doc.add_heading('Why We Chose It', level=2)
doc.add_paragraph('Focuses on the hardest-to-classify counties (the "Moderate" and "Volatile" middle classes where RF struggles)', style='List Bullet')
doc.add_paragraph('Has built-in protection against overfitting (regularization penalties)', style='List Bullet')
doc.add_paragraph('Integrates with SHAP \u2014 a powerful tool for explaining individual predictions', style='List Bullet')
doc.add_paragraph('Consistently the top-performing model on tabular (spreadsheet-like) data in research competitions', style='List Bullet')

doc.add_heading('Assumptions', level=2)
doc.add_paragraph('Each iteration meaningfully improves on the last (learning rate controls speed)', style='List Bullet')
doc.add_paragraph('Regularization prevents the model from memorizing noise in our small dataset', style='List Bullet')

doc.add_heading('Hyperparameter Tuning', level=2)
add_table(doc,
    ['Setting', 'Options Tested', 'Best', 'What It Controls'],
    [
        ['Number of trees', '100\u2013500', '300', 'Total learning capacity'],
        ['Max depth', '3\u20138', '5', 'Complexity per tree'],
        ['Learning rate', '0.01\u20130.2', '0.1', 'How much each tree contributes'],
        ['Subsample', '60\u2013100%', '80%', 'Fraction of data per tree'],
        ['Column sample', '60\u2013100%', '80%', 'Fraction of features per tree'],
        ['L1 regularization', '0\u20131', '0.1', 'Penalizes unnecessary features'],
        ['L2 regularization', '1\u201310', '5', 'Penalizes large predictions'],
    ]
)

doc.add_heading('5-Class Results', level=2)
add_table(doc,
    ['Metric', 'Value', 'vs Random Forest'],
    [
        ['Accuracy', '0.388', '+0.020 (better)'],
        ['Precision (macro)', '0.396', '+0.027 (better)'],
        ['Recall (macro)', '0.383', '+0.015 (better)'],
        ['F1-macro', '0.387', '+0.025 (better)'],
        ['ROC-AUC', '0.783', '\u22120.020 (slightly worse)'],
    ]
)
doc.add_paragraph(
    'XGBoost achieves the best overall 5-class performance (F1 = 0.387). '
    'It particularly outperforms Random Forest on the "Stable" middle class (F1 = 0.37 vs 0.18), '
    'confirming that the "learning from mistakes" approach helps with the hardest cases.'
)

add_figure(doc, 'xgb_confusion_matrix.png',
           'Figure 4: XGBoost confusion matrix. Better middle-class separation than RF, '
           'particularly for the "Stable" (Q2) class.')

doc.add_heading('SHAP Analysis \u2014 Explaining Individual Predictions', level=2)

add_method_explainer(doc, 'SHAP (SHapley Additive exPlanations)',
    'SHAP is a method for explaining why the model made a specific prediction for a specific county. '
    'For each county, SHAP calculates how much each demographic feature "pushed" the prediction '
    'toward high or low volatility. Think of it as an itemized receipt for each prediction: '
    '"This county was predicted Highly Volatile because: +0.85 from high rent, +0.40 from high diversity, '
    '\u22120.30 from high homeownership." It comes from game theory and provides mathematically rigorous '
    'explanations.')

add_figure(doc, 'xgb_shap_summary.png',
           'Figure 5: SHAP summary \u2014 each dot is one county. Color indicates feature value '
           '(red = high, blue = low). Features ranked by importance. Housing cost and racial diversity dominate.')

add_figure(doc, 'xgb_shap_dependence_top5.png',
           'Figure 6: How the top 5 features affect predictions. Each dot is a county. '
           'The diversity index (race_entropy_norm) shows a clear "step" \u2014 volatility jumps once '
           'diversity crosses a threshold.')

doc.add_heading('Feature Interactions \u2014 Demographics Don\'t Act Alone', level=2)
doc.add_paragraph(
    'One of the most important findings is that demographic features interact with each other. '
    'Racial diversity alone doesn\'t predict volatility \u2014 it\'s diversity combined with '
    'urbanization, education, or immigration that creates instability.'
)

add_table(doc,
    ['Feature Pair', 'Strength', 'Plain-Language Meaning'],
    [
        ['Diversity \u00d7 Population Density', '0.133', 'Diversity predicts volatility\nmost in suburban/urban areas'],
        ['Education \u00d7 Diversity', '0.091', 'Educated + diverse counties\nare especially volatile'],
        ['Population \u00d7 Density', '0.078', 'Captures the urban-suburban-\nrural gradient'],
        ['Diversity \u00d7 Population Size', '0.069', 'Diversity matters more\nin large counties'],
        ['HS Education \u00d7 Diversity', '0.057', 'Lower-education diverse areas\nare especially volatile'],
        ['Immigration \u00d7 Diversity', '0.052', 'Immigration-driven diversity\nhas an extra volatility signal'],
    ]
)

bold_para(doc, 'The diversity index appears in 6 of the top 10 interactions.',
          ' It is the "hub" feature \u2014 its signal is amplified or dampened depending on whether '
          'the county is urban or rural, educated or not, immigrant-heavy or native-born.')

add_figure(doc, 'deepdive_shap_interactions.png',
           'Figure 7: Top 6 SHAP interaction scatter plots. Each dot is a county; '
           'color shows the value of the interacting feature.')

add_figure(doc, 'deepdive_interaction_heatmap.png',
           'Figure 8: Interaction strength heatmap. Brighter = stronger interaction. '
           'The diversity index row and column are the brightest \u2014 it interacts with everything.')

doc.add_heading('Challenges & Solutions', level=2)
doc.add_paragraph('Challenge: XGBoost requires class labels starting at 0 (ours start at 1). Solution: Subtract 1 before training, add 1 back after prediction.', style='List Bullet')
doc.add_paragraph('Challenge: Risk of overfitting with sequential learning. Solution: Strong regularization (L1=0.1, L2=5) plus only using 80% of data/features per tree.', style='List Bullet')

add_takeaway_box(doc,
    'XGBoost is the best model for the detailed 5-class task. Its SHAP analysis reveals that '
    'demographics don\'t act alone \u2014 it\'s the combination of racial diversity with urbanization, '
    'education, and immigration that creates electoral instability. '
    'The diversity index is the "hub" that connects to every other important feature.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. MODEL C — SVM
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('6. Model C \u2014 SVM (Support Vector Machine)', level=1)

add_method_explainer(doc, 'SVM (Support Vector Machine)',
    'SVM draws a dividing line (or "hyperplane") between classes of counties and tries '
    'to make that line as far as possible from the nearest counties on each side. '
    'Think of sorting balls on a table by color: SVM finds the widest possible gap between groups. '
    'The "support vectors" are the balls closest to the gap \u2014 they define where the boundary sits. '
    'The "kernel trick" allows SVM to bend the dividing surface into curves (RBF kernel) '
    'rather than just straight lines (Linear kernel).')

doc.add_heading('Why We Chose It', level=2)
doc.add_paragraph('Focuses on the hardest-to-classify boundary cases (support vectors), not all data', style='List Bullet')
doc.add_paragraph('Testing Linear vs curved (RBF) kernel tells us whether the demographic-volatility relationship is simple or complex', style='List Bullet')
doc.add_paragraph('Provides theoretical guarantees about generalization to new data', style='List Bullet')

doc.add_heading('Assumptions', level=2)
doc.add_paragraph('Linear kernel: The boundary between volatile and stable counties is approximately a straight line in demographic space', style='List Bullet')
doc.add_paragraph('RBF kernel: Counties that are similar demographically should behave similarly electorally', style='List Bullet')
doc.add_paragraph('Features must be on the same scale (requires StandardScaler)', style='List Bullet')

doc.add_heading('Hyperparameter Tuning', level=2)
add_table(doc,
    ['Kernel', 'Setting', 'Options', 'Best'],
    [
        ['Linear', 'C (penalty)', '0.01\u2013100', '1'],
        ['Linear', 'Class weight', 'balanced / none', 'balanced'],
        ['RBF', 'C (penalty)', '0.01\u2013100', '10'],
        ['RBF', 'Gamma (curve tightness)', 'scale, auto, 0.01\u20131', 'scale'],
        ['RBF', 'Class weight', 'balanced / none', 'balanced'],
    ]
)

doc.add_heading('Kernel Comparison \u2014 A Diagnostic Finding', level=2)
add_table(doc,
    ['Metric', 'Linear SVM', 'RBF (Curved) SVM'],
    [
        ['Accuracy', '0.368', '0.340'],
        ['F1-macro', '0.364', '0.347'],
        ['ROC-AUC', '0.765', '0.741'],
    ]
)

doc.add_paragraph(
    'In the 5-class task, the straight-line (Linear) kernel beats the curved (RBF) kernel. '
    'This tells us something important about the data: when trying to distinguish all 5 levels '
    'of volatility, the boundaries between them are approximately straight lines in demographic space.'
)
bold_para(doc, 'But this reverses in binary classification:',
          ' When we simplify to just High vs Low volatility, the curved kernel wins '
          '(RBF AUC=0.789 vs Linear AUC=0.734). This means the single boundary between '
          '"volatile" and "stable" counties is moderately curved \u2014 consistent with the '
          'interaction effects we found in SHAP analysis (diversity matters differently '
          'depending on urbanization and education).')

add_figure(doc, 'svm_confusion_matrix.png',
           'Figure 9: SVM confusion matrix. Similar pattern to other models: '
           'strong at extremes, weaker in the middle.')

doc.add_heading('Challenges & Solutions', level=2)
doc.add_paragraph('Challenge: SVM provides no built-in feature importance. Solution: Used RF and XGBoost for importance rankings; SVM served as a performance benchmark and diagnostic tool (kernel comparison).', style='List Bullet')
doc.add_paragraph('Challenge: SVM is sensitive to the regularization parameter C. Solution: Nested cross-validation searches over multiple C values independently.', style='List Bullet')

add_takeaway_box(doc,
    'SVM\'s main contribution is diagnostic: the kernel comparison reveals that the '
    'volatile/stable boundary is moderately nonlinear. This confirms the interaction effects \u2014 '
    'demographics interact in non-additive ways to produce electoral instability. '
    'A straight line cannot fully separate volatile from stable counties.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. MODEL D — MISFIT DETECTOR
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('7. Model D \u2014 Demographic Misfit Detector (Logistic Regression)', level=1)

add_method_explainer(doc, 'Logistic Regression',
    'Logistic Regression is one of the simplest classification models. '
    'It draws a straight line through the demographic data to separate Democratic-leaning '
    'counties from Republican-leaning ones, and outputs a probability (e.g., "this county has '
    'a 73% chance of voting Democrat based on its demographics"). It\'s interpretable and fast, '
    'but assumes the relationship is approximately linear.')

doc.add_heading('A Creative Two-Stage Approach', level=2)
doc.add_paragraph(
    'Rather than directly predicting volatility, we use Logistic Regression in a creative way:'
)
doc.add_paragraph(
    'Stage 1: Predict which party a county should vote for, based purely on its demographics.',
    style='List Bullet')
doc.add_paragraph(
    'Stage 2: Find "misfits" \u2014 counties where the prediction is confidently wrong. '
    'If demographics say a county should vote Republican with 96% confidence, but it actually '
    'votes Democrat, that county is a "demographic misfit." We hypothesize that misfits are volatile.',
    style='List Bullet')

doc.add_paragraph(
    'Why? Because a county that defies its demographic expectations is likely undergoing some kind '
    'of transition \u2014 demographic change, political realignment, or unique local dynamics \u2014 '
    'that should also make its elections less predictable.'
)

doc.add_heading('Performance', level=2)
bold_para(doc, 'Partisan prediction accuracy: 89.6%',
          ' \u2014 demographics alone correctly predict which party wins in 9 out of 10 counties. '
          'The 10% that the model gets wrong are precisely the interesting cases.')

doc.add_heading('Top 5 Demographic Misfits \u2014 Counties That Defy Their Demographics', level=2)
add_table(doc,
    ['County', 'State', 'Model Says', 'Actually Votes', 'Misfit Score', 'Volatility'],
    [
        ['Leelanau', 'MI', '96% Rep', 'Dem (+8%)', '0.96', 'Highly Volatile'],
        ['Marquette', 'MI', '96% Rep', 'Dem (+9%)', '0.96', 'Highly Volatile'],
        ['Nash', 'NC', '95% Dem', 'Rep (\u22122%)', '0.95', 'Very Stable'],
        ['Lenoir', 'NC', '94% Dem', 'Rep (\u22127%)', '0.94', 'Stable'],
        ['Cabarrus', 'NC', '87% Dem', 'Rep (\u22128%)', '0.87', 'Highly Volatile'],
    ]
)

doc.add_paragraph(
    'Leelanau and Marquette are small, rural, predominantly white Michigan counties \u2014 '
    'demographics that scream "Republican." Yet they vote solidly Democrat, driven by university '
    'communities and resort/tourism economies. They are also Highly Volatile, '
    'confirming that defying demographic expectations correlates with electoral instability.'
)

doc.add_heading('The Misfit-Volatility Link', level=2)
add_table(doc,
    ['Finding', 'Statistical Test', 'Value'],
    [
        ['Misfits are more volatile', 'Spearman \u03c1', '0.408 (p < 0.001)'],
        ['Misfits don\'t lean one direction', 'Pearson r', '\u22120.017 (no relationship)'],
    ]
)
doc.add_paragraph(
    'In plain terms: counties that defy demographic expectations are significantly more volatile, '
    'but being a misfit does not predict whether they lean left or right. '
    'Misfits are unstable in both directions.'
)

add_figure(doc, 'misfit_scatter.png',
           'Figure 10: Each county plotted by predicted partisanship (x-axis) vs actual margin (y-axis). '
           'Dot size = volatility. Counties far from the diagonal are "misfits."')

add_figure(doc, 'misfit_correlation_matrix.png',
           'Figure 11: Cross-correlations between misfit score, volatility, swing direction, and P(Dem). '
           'Misfit and volatility are clearly linked (\u03c1 = 0.408).')

doc.add_heading('Who Are the Misfits? \u2014 A Demographic Profile', level=2)
add_table(doc,
    ['Characteristic', 'Misfits', 'Non-Misfits', 'p-value'],
    [
        ['Racial diversity', '0.567', '0.419', '< 0.0001'],
        ["Bachelor's degree+", '30.9%', '24.7%', '< 0.0001'],
        ['Homeownership', '71.1%', '76.3%', '< 0.0001'],
        ['Median rent', '$1,056', '$932', '< 0.0001'],
        ['Median age', '42.5 yrs', '44.5 yrs', '0.003'],
    ]
)
doc.add_paragraph(
    'Misfits are younger, more diverse, better educated, higher-rent, and less rooted '
    '(lower homeownership). In a word: they are communities "in flux" \u2014 '
    'demographically transitioning places where old voting patterns no longer hold.'
)

doc.add_heading('Two Types of Misfits', level=2)
add_table(doc,
    ['Type', 'Count', 'Avg Volatility', 'Where'],
    [
        ['"Surprise Dem"\n(expected Rep, votes Dem)', '8 of 30', '0.807 (very high)', 'MI university/resort towns,\nPA Scranton area'],
        ['"Surprise Rep"\n(expected Dem, votes Rep)', '22 of 30', '0.383 (moderate)', 'NC rural South'],
    ]
)

bold_para(doc, '"Surprise Democrat" misfits are twice as volatile as "Surprise Republican" ones.',
          ' University towns and resort communities that vote against their rural/white demographics '
          'are the most electorally unstable counties in the entire dataset. '
          'For campaign strategists, these are high-priority targets with outsized swing potential.')

add_figure(doc, 'deepdive_misfit_by_volatility.png',
           'Figure 12: Misfit scores by volatility class. The proportion of misfits rises monotonically '
           'from 6% (Very Stable) to 34% (Highly Volatile). Kruskal-Wallis p < 0.000001.')

add_takeaway_box(doc,
    'The Misfit Detector reveals that counties whose voting behavior defies their demographics '
    'tend to be more electorally volatile. These "demographic misfits" are younger, more diverse, '
    'more educated, and more transient \u2014 communities in demographic transition. '
    '"Surprise Democrat" misfits (rural areas voting blue) are the most volatile of all.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 8. BINARY CLASSIFICATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('8. Binary Classification \u2014 High vs Low Volatility', level=1)

doc.add_heading('Why Simplify to Binary?', level=2)
doc.add_paragraph(
    'The 5-class task (F1 ~ 0.39) is limited because the middle classes ("Stable", "Moderate", '
    '"Volatile") overlap heavily in demographic space \u2014 it\'s hard to tell them apart. '
    'But campaign strategists don\'t need to distinguish between "Moderate" and "Volatile." '
    'They need to answer one question: "Will this county swing significantly or not?" '
    'Collapsing to binary \u2014 High (top 40%) vs Low (bottom 60%) \u2014 dramatically improves performance.'
)

doc.add_heading('Results \u2014 All Models Compared', level=2)
add_table(doc,
    ['Model', 'Accuracy', 'F1', 'Precision', 'Recall', 'AUC'],
    [
        ['Random Forest', '0.780', '0.718', '0.726', '0.690', '0.828'],
        ['XGBoost', '0.748', '0.667', '0.718', '0.610', '0.803'],
        ['SVM (RBF)', '0.760', '0.674', '0.724', '0.630', '0.789'],
        ['Logistic Regression', '0.680', '0.604', '0.598', '0.610', '0.745'],
        ['SVM (Linear)', '0.688', '0.610', '0.610', '0.610', '0.734'],
    ]
)

doc.add_paragraph(
    'In plain terms: the Random Forest correctly identifies 78% of counties as high or low volatility, '
    'and when it says a county is high-volatility, it\'s right 73% of the time (precision). '
    'It catches 69% of all truly high-volatility counties (recall).'
)

doc.add_heading('The Power of Problem Framing', level=2)
add_table(doc,
    ['Metric', '5-Class Best', 'Binary Best', 'Improvement'],
    [
        ['F1', '0.387 (XGBoost)', '0.718 (RF)', '+85%'],
        ['AUC', '0.803 (RF)', '0.828 (RF)', '+3%'],
    ]
)
bold_para(doc, 'Simplifying from 5 classes to 2 nearly doubles our accuracy.',
          ' This is one of the most important lessons: choosing the right question to ask '
          '(binary vs 5-class) matters more than choosing the right algorithm.')

add_figure(doc, 'deepdive_binary_roc_all_models.png',
           'Figure 13: ROC curves for all 5 binary models. The curve shows the trade-off between '
           'catching true positives and avoiding false alarms. Random Forest (blue) is closest to the top-left corner.')

add_figure(doc, 'deepdive_binary_confusion_matrices.png',
           'Figure 14: Confusion matrices for all 5 binary models side-by-side.')

add_figure(doc, 'deepdive_shap_binary_summary.png',
           'Figure 15: SHAP summary for binary classification. Diversity and housing cost dominate.')

doc.add_heading('How Confident Are We? \u2014 Bootstrap Analysis', level=2)
doc.add_paragraph(
    'To ensure our results aren\'t a fluke of one particular data split, '
    'we repeated the entire evaluation 200 times with randomly resampled data:'
)
add_table(doc,
    ['Model', 'F1 [95% CI]', 'AUC [95% CI]', 'Accuracy [95% CI]'],
    [
        ['Random Forest', '0.818 [0.764, 0.872]', '0.918 [0.882, 0.946]', '0.856 [0.816, 0.900]'],
        ['XGBoost', '0.845 [0.789, 0.899]', '0.932 [0.894, 0.966]', '0.880 [0.840, 0.920]'],
    ]
)
doc.add_paragraph(
    'The confidence intervals are tight, confirming that our performance estimates are robust '
    'and not an artifact of a lucky data split.'
)

add_figure(doc, 'deepdive2_feature_ablation.png',
           'Figure 16: Left: What happens when we remove each feature group? '
           'Right: How important is each individual feature (with 95% confidence bars)?')

add_takeaway_box(doc,
    'Binary classification achieves F1 = 0.718 \u2014 nearly twice the 5-class performance. '
    'Random Forest is the clear winner. The lesson: asking the right question ("will it swing?") '
    'matters more than choosing the right algorithm. Bootstrap analysis confirms these results are robust.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 9. ENSEMBLE STACKING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('9. Ensemble Stacking \u2014 Can We Do Better by Combining Models?', level=1)

add_method_explainer(doc, 'Ensemble Stacking',
    'Ensemble stacking combines multiple models together. "Soft Voting" averages the predictions '
    'from RF, XGBoost, and SVM. "Stacking" goes further: it trains a second-level model '
    '(Logistic Regression) to learn the optimal way to combine the first-level predictions. '
    'The idea is that different models catch different patterns, so combining them should improve accuracy.')

doc.add_heading('Results', level=2)
add_table(doc,
    ['Model', 'F1', 'AUC', 'Type'],
    [
        ['Random Forest', '0.718', '0.828', 'Individual (best)'],
        ['Soft Voting (RF+XGB+SVM)', '0.691', '0.826', 'Ensemble'],
        ['XGBoost', '0.667', '0.803', 'Individual'],
        ['SVM (RBF)', '0.663', '0.819', 'Individual'],
        ['Stacking (LR meta)', '0.647', '0.816', 'Ensemble'],
        ['Logistic Regression', '0.604', '0.745', 'Individual'],
    ]
)

add_figure(doc, 'deepdive2_model_scoreboard.png',
           'Figure 17: Final scoreboard. Left: F1 scores. Right: ROC curves. '
           'Random Forest alone outperforms all ensemble combinations.')

doc.add_paragraph(
    'Surprisingly, combining models does NOT improve performance. Stacking actually hurts '
    '(F1 drops from 0.718 to 0.647). Why? With only 250 counties, there isn\'t enough data '
    'for the second-level model to learn a useful combination \u2014 it "oversmooths" '
    'the predictions and loses the nuances that Random Forest captures on its own.'
)

add_takeaway_box(doc,
    'More complex is not always better. Ensembles fail to improve on Random Forest because '
    'our dataset (250 counties) is too small for the stacking approach to learn anything useful. '
    'The performance ceiling is determined by dataset size, not by algorithm sophistication. '
    'This is an important practical finding: when your data is small, a well-tuned simple model beats a complex one.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. MODEL COMPARISON
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('10. Model Comparison & Performance Evaluation', level=1)

doc.add_heading('10.1 Complete 5-Class Comparison', level=2)
add_table(doc,
    ['Model', 'Accuracy', 'F1-macro', 'ROC-AUC', 'Best At'],
    [
        ['XGBoost', '0.388', '0.387', '0.783', 'Best overall 5-class'],
        ['SVM (Linear)', '0.368', '0.364', '0.765', 'Best at middle classes'],
        ['Random Forest', '0.368', '0.362', '0.803', 'Best ROC-AUC'],
    ]
)

add_figure(doc, 'per_class_f1_heatmap.png',
           'Figure 18: Per-class F1 heatmap. All models struggle with middle classes '
           'but succeed at identifying the extremes.')

doc.add_heading('10.2 Complete Binary Comparison', level=2)
add_table(doc,
    ['Model', 'Accuracy', 'F1', 'AUC', 'Type'],
    [
        ['Random Forest', '0.780', '0.718', '0.828', 'Best overall'],
        ['Soft Voting', '0.768', '0.691', '0.826', 'Ensemble'],
        ['XGBoost', '0.748', '0.667', '0.803', 'Individual'],
        ['SVM (RBF)', '0.760', '0.663', '0.819', 'Individual'],
        ['Stacking', '0.760', '0.647', '0.816', 'Ensemble'],
        ['Logistic Regression', '0.680', '0.604', '0.745', 'Individual'],
    ]
)

doc.add_heading('10.3 Why Random Forest Wins (Binary)', level=2)
doc.add_paragraph('With only 250 counties, the variance-reducing power of bagging (averaging many independent trees) matters more than boosting\'s ability to correct errors sequentially.', style='List Bullet')
doc.add_paragraph('Random Forest naturally captures the feature interactions we identified (diversity \u00d7 urbanization, education \u00d7 density) without needing them to be explicitly specified.', style='List Bullet')
doc.add_paragraph('It is less sensitive to hyperparameter choices than XGBoost \u2014 a practical advantage when data is scarce.', style='List Bullet')

doc.add_heading('10.4 Feature Group Ablation \u2014 Which Demographic Themes Matter?', level=2)
doc.add_paragraph(
    'We systematically removed each group of features to measure its contribution:'
)
add_table(doc,
    ['Feature Group', 'Features', 'F1 without it', 'F1 drop', 'F1 with it alone'],
    [
        ['Race / Ethnicity', '6', '0.626', '+0.092 (biggest)', '0.619 (86% of full)'],
        ['Urbanization', '6', '0.684', '+0.034', '0.583'],
        ['Education', '2', '0.708', '+0.010', '0.583'],
        ['Housing', '3', '0.712', '+0.006', '0.612 (85% alone)'],
        ['Economic', '5', '0.722', '\u22120.004', '0.570'],
        ['Household / Age', '6', '0.728', '\u22120.010', '0.471'],
    ]
)

doc.add_paragraph(
    'The Race/Ethnicity group is both the most necessary (removing it causes the biggest drop) '
    'and the most sufficient (using it alone achieves 86% of full performance). '
    'The Economic and Household/Age groups are actually dispensable \u2014 '
    'removing them slightly improves performance, suggesting they add noise.'
)

add_takeaway_box(doc,
    'Random Forest is the best model overall (binary F1 = 0.718). XGBoost wins the 5-class task. '
    'Race/Ethnicity features are the most important group by a wide margin \u2014 they alone '
    'achieve 86% of the full model\'s performance. Economic variables and age/household structure '
    'add minimal value. The key predictors are diversity, housing cost, poverty, and education.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 11. CROSS-STATE GENERALIZATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('11. Cross-State Generalization Experiment', level=1)

doc.add_heading('The Key Question', level=2)
doc.add_paragraph(
    'If we build a model using Pennsylvania and North Carolina counties, '
    'can it predict which Michigan counties will be volatile? '
    'In other words, do the same demographics predict volatility everywhere, '
    'or is politics fundamentally local?'
)

doc.add_heading('Setup', level=2)
doc.add_paragraph(
    'Leave-One-State-Out (LOSO): Train on two states, test on the third. '
    'We use the best 5-class model (XGBoost). '
    'Critically, we refit the data scaler on training states only \u2014 '
    'the test state\'s data is not used in any way during training.'
)

doc.add_heading('Results \u2014 A Striking Failure', level=2)
add_table(doc,
    ['Test State', 'Train States', 'Accuracy', 'F1-macro'],
    [
        ['Michigan', 'PA + NC', '0.205', '0.182'],
        ['North Carolina', 'PA + MI', '0.210', '0.213'],
        ['Pennsylvania', 'MI + NC', '0.209', '0.189'],
        ['', 'Average', '0.208', '0.195'],
        ['', 'Random guessing (1/5)', '0.200', '0.200'],
    ]
)

bold_para(doc, 'The model trained on two states performs no better than random guessing on the third.',
          ' This is one of the most important findings in the entire analysis.')

add_figure(doc, 'loso_confusion_matrices.png',
           'Figure 19: LOSO confusion matrices. All three rotations collapse to near-random predictions \u2014 '
           'the model simply cannot transfer what it learned from one state to another.')

doc.add_heading('Why Does This Happen? \u2014 Each State Has a Different Volatility Signature', level=2)
doc.add_paragraph(
    'To understand this failure, we trained separate models within each state '
    'and compared which features matter most:'
)

add_table(doc,
    ['Rank', 'Pennsylvania', 'Michigan', 'North Carolina'],
    [
        ['#1', 'Racial diversity', 'Education (bachelor\'s+)', 'Poverty rate'],
        ['#2', 'Housing cost (rent)', 'Housing cost (rent)', 'Housing cost (rent)'],
        ['#3', 'Immigration', 'Work from home', 'Population density'],
        ['#4', 'Multiracial pop.', 'Home values', 'Home values'],
        ['#5', 'Income', 'Senior population', 'Education'],
    ]
)

bold_para(doc, 'The only feature that matters in all three states: housing cost.',
          ' Everything else is state-specific. Pennsylvania\'s volatility is driven by racial diversity, '
          'Michigan\'s by education levels, and North Carolina\'s by poverty. '
          'These reflect fundamentally different political cultures.')

add_figure(doc, 'deepdive_state_feature_importance.png',
           'Figure 20: Per-state feature importance (3 panels). The feature rankings are dramatically '
           'different in each state, explaining why cross-state prediction fails.')

add_figure(doc, 'deepdive_state_roc_curves.png',
           'Figure 21: State-specific ROC curves. Within-state models work well (AUC 0.78\u20130.84); '
           'cross-state models fail completely.')

add_table(doc,
    ['State', 'Counties', 'Within-State AUC', 'Within-State F1'],
    [
        ['Pennsylvania', '67', '0.838', '0.591'],
        ['Michigan', '83', '0.778', '0.636'],
        ['North Carolina', '100', '0.829', '0.736'],
    ]
)

add_takeaway_box(doc,
    'Demographics predict volatility well within each state (AUC 0.78\u20130.84) but fail completely '
    'across state lines (F1 drops to random baseline). The same demographic profile produces different '
    'political behavior in different states. This means national-level models are fundamentally '
    'limited \u2014 effective prediction requires state-specific analysis. '
    'Housing cost is the only universal predictor.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 12. HYPOTHESIS TESTING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('12. Hypothesis Testing', level=1)

doc.add_heading('H1: Volatility Inversely Correlated with Wealth \u2014 CONFIRMED', level=2)
add_table(doc,
    ['Evidence Source', 'Method', 'Result'],
    [
        ['SHAP analysis', 'Income SHAP vs Q5', 'r = \u22120.602 (strong negative, p < 0.0001)'],
        ['Feature importance', 'Poverty rate ranking', 'Top 5 in all models'],
        ['Partial dependence', 'Rent effect shape', 'Monotonic increase: higher rent \u2192 higher P(volatile)'],
        ['Feature ablation', 'Economic group alone', 'F1 = 0.570 (works but insufficient alone)'],
    ]
)
doc.add_paragraph(
    'In plain terms: wealthier counties are indeed more stable. The SHAP analysis shows a strong '
    'negative correlation (r = \u22120.602) between income and high-volatility prediction. '
    'However, income alone is not enough to predict volatility \u2014 it matters most in combination '
    'with diversity and urbanization. A wealthy, diverse suburb is more volatile than a wealthy, '
    'homogeneous exurb.'
)

doc.add_heading('H2: Diversity Index > Individual Race Features \u2014 CONFIRMED', level=2)
add_table(doc,
    ['Evidence Source', 'Method', 'Result'],
    [
        ['Permutation importance', 'RF, 30 repeats', 'race_entropy_norm = #1 (top of all features)'],
        ['SHAP importance', 'XGBoost mean |SHAP|', 'Entropy (0.397) > pct_black (0.347)'],
        ['Partial dependence', 'Probability swing range', 'Entropy = 0.14 (largest of all features)'],
        ['Feature ablation', 'Race group alone', 'F1 = 0.619 (86% of baseline with just 6 race features)'],
        ['Interaction analysis', 'SHAP interactions', 'Entropy in 6 of top 10 interactions'],
        ['Threshold analysis', 'Maximum-separation + bootstrap', 'Odds ratio = 7.15\u00d7 above threshold'],
    ]
)
doc.add_paragraph(
    'Confirmed across 6 independent lines of evidence. The racial diversity index outperforms every '
    'individual race variable as a predictor. This means it\'s not about the presence of any particular '
    'racial group \u2014 it\'s about how mixed the county is. Counties where no single group dominates '
    '(entropy > 0.55) are 7 times more likely to be highly volatile. '
    'This makes intuitive sense: racially homogeneous communities have more settled political identities, '
    'while mixed communities are still negotiating theirs.'
)

add_takeaway_box(doc,
    'Both hypotheses confirmed. H1: Wealth correlates with stability (r = \u22120.602). '
    'H2: The diversity index outperforms all individual race variables \u2014 it\'s racial mixing, '
    'not any single group, that predicts volatility. Counties above the diversity threshold '
    'are 7\u00d7 more likely to swing.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 13. ADVANCED ANALYSIS — DEEP-DIVES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('13. Advanced Analysis \u2014 Deep-Dives', level=1)

doc.add_heading('13.1 Temporal Demographic Shifts (2016 \u2192 2024)', level=2)
doc.add_paragraph(
    'Beyond static demographics, we asked: does demographic change over time predict volatility? '
    'We computed how each county\'s demographics shifted between 2016 and 2024 census data '
    'and correlated these changes with electoral volatility.'
)

add_table(doc,
    ['Demographic Change', 'Correlation', 'p-value', 'Meaning'],
    [
        ['Growing Hispanic pop.', '+0.220', '0.0005', 'Counties gaining Hispanic residents swing more'],
        ['Rising rents', '+0.209', '0.0009', 'Gentrifying areas are volatile'],
        ['Growing multiracial pop.', '+0.209', '0.0009', 'Increasing racial mixing \u2192 more swings'],
        ['Growing Black pop.', '\u22120.194', '0.0020', 'Growth of one group \u2192 less volatile'],
    ]
)

doc.add_paragraph(
    'The pattern is consistent with our entropy finding: diversification (more mixing) predicts '
    'volatility, while growth of a single group does not. Counties becoming more racially mixed '
    'over time are the ones whose politics are most in flux.'
)

add_figure(doc, 'deepdive2_temporal_demographic_shifts.png',
           'Figure 22: Temporal shifts vs volatility. Left: correlation strengths. Right: scatter plots '
           'for the top two predictors (Hispanic growth and rent increases).')

add_takeaway_box(doc,
    'It\'s not just where demographics are \u2014 it\'s where they\'re going. '
    'Counties becoming more racially mixed and experiencing rising housing costs '
    'are the most volatile. This suggests volatility is driven by demographic transition, not static profiles.')

doc.add_heading('13.2 Election Margin Trajectories (2016 \u2192 2020 \u2192 2024)', level=2)
doc.add_paragraph(
    'Instead of reducing three elections to a single volatility number, '
    'we classified each county by how its margin evolved across all three cycles:'
)

add_table(doc,
    ['Trajectory', 'Counties', 'Mean Vol', '% High Vol', 'What Happened'],
    [
        ['Blue Bounceback', '117 (47%)', '\u22120.17', '34%', 'Swung Democrat in 2020,\nthen snapped back in 2024'],
        ['Steady Red Drift', '94 (38%)', '+0.39', '50%', 'Moved Republican in both\ncycles \u2014 most volatile'],
        ['Steady Blue Drift', '35 (14%)', '\u22120.28', '37%', 'Moved Democrat in both cycles'],
        ['Red Bounceback', '4 (2%)', '\u22121.82', '0%', 'Moved R then D \u2014 extremely rare'],
    ]
)

doc.add_paragraph(
    'The dominant pattern is "Blue Bounceback" (47% of counties) \u2014 the 2020 Biden surge '
    'was temporary in most places, with margins reverting toward pre-2020 levels in 2024. '
    'But the most volatile trajectory is "Steady Red Drift" (50% high volatility) \u2014 '
    'counties that moved Republican in both 2020 and 2024 are the most unstable.'
)

bold_para(doc, 'North Carolina is different:',
          ' 49% of NC counties show Steady Red Drift vs only 27% (PA) and 33% (MI). '
          'NC is undergoing a more profound, ongoing political realignment than the northern states.')

add_figure(doc, 'deepdive2_margin_trajectories.png',
           'Figure 23: Margin trajectories by state. Each line is one county. Colors show trajectory type. '
           'The "spaghetti" of lines illustrates the diversity of paths counties take.')

add_takeaway_box(doc,
    'The 2020 Biden surge was temporary in most counties ("Blue Bounceback" = 47%). '
    'Counties drifting steadily Republican are the most volatile (50% high vol). '
    'NC is experiencing a deeper realignment than PA or MI.')

doc.add_heading('13.3 The Diversity Threshold \u2014 A Precise Tipping Point', level=2)
doc.add_paragraph(
    'Our analysis consistently shows that racial diversity acts like a "switch" rather than a gradual '
    'dial. We estimated the precise threshold using a statistical optimization technique '
    'with 1000 bootstrap resamples for confidence:'
)

add_table(doc,
    ['Metric', 'Value', 'Meaning'],
    [
        ['Optimal threshold', '0.552', 'No single race exceeds ~60% of population'],
        ['95% Confidence Interval', '[0.397 \u2013 0.642]', 'The switch zone'],
        ['Below threshold', '24.4% high vol', 'About 1 in 4 volatile'],
        ['Above threshold', '69.8% high vol', 'About 7 in 10 volatile'],
        ['Odds ratio', '7.15\u00d7', 'Mixed counties are 7\u00d7 more likely to swing'],
    ]
)

doc.add_paragraph(
    'In practical terms: a county where no single racial group makes up more than ~60% of the '
    'population is 7 times more likely to be highly volatile. This is not a gradual relationship \u2014 '
    'the probability of high volatility jumps from about 20% to about 70% as you cross this threshold. '
    'Think of it as a "tipping point" for political instability.'
)

add_figure(doc, 'deepdive2_diversity_threshold.png',
           'Figure 24: The diversity threshold. Left: overlapping distributions for high vs low volatility counties. '
           'Right: the probability of high volatility jumps from ~20% to ~70% across the threshold zone.')

add_takeaway_box(doc,
    'There is a clear tipping point at race_entropy_norm \u2248 0.55. Below it, counties are usually '
    'stable. Above it, they are 7\u00d7 more likely to swing. This is not gradual \u2014 it\'s a switch. '
    'Counties where no single racial group exceeds ~60% are in the "volatile zone."')

doc.add_heading('13.4 Partial Dependence \u2014 True Shape of Feature Effects', level=2)

add_method_explainer(doc, 'Partial Dependence Plots (PDPs)',
    'A PDP shows how one feature affects the model\'s prediction, averaged across all other features. '
    'Imagine taking all 250 counties and asking: "If we could magically change just this one feature '
    '(say, rent) while keeping everything else the same, how would the predicted volatility change?" '
    'The resulting curve reveals the true functional shape of each feature\'s effect.')

add_table(doc,
    ['Feature', 'Shape', 'What It Means'],
    [
        ['Diversity index', 'Step function', 'Flat \u2192 jumps at threshold \u2192 plateaus.\nConfirms the "switch" finding.'],
        ['Housing cost (rent)', 'Steady increase', 'Higher rent = steadily higher\nvolatility probability'],
        ['Poverty rate', 'Late spike', 'Flat for most values, then jumps\nat extreme poverty'],
        ['Married-couple %', 'Flat', 'Works through interactions only,\nnot on its own'],
        ['Bachelor\'s degree %', 'Slight increase', 'Modest effect at high education'],
        ['Population density', 'Uptick at extremes', 'Mostly flat; works through\ninteractions'],
    ]
)

add_figure(doc, 'deepdive2_partial_dependence.png',
           'Figure 25: Partial dependence plots (top 6 features), 2D interaction contours, '
           'and PDP-based feature importance. The diversity "step function" is clearly visible.')

add_takeaway_box(doc,
    'PDPs reveal that diversity operates as a switch (step function), rent as a steady escalator, '
    'and poverty as a late spike. Several important features (married couples, density) have no '
    'individual effect \u2014 they work only through interactions with other features.')

doc.add_heading('13.5 Nonlinear Deep-Dives \u2014 Feature Interactions in Plain Language', level=2)

bold_para(doc, 'Diversity \u00d7 Poverty \u2014 The Clearest Finding:')
add_table(doc,
    ['Quadrant', 'Counties', 'Mean Vol', '% High Vol', '% Misfits'],
    [
        ['Low Diversity + Low Poverty', '73', '\u22120.622', '19%', '11%'],
        ['Low Diversity + High Poverty', '52', '\u22120.669', '23%', '8%'],
        ['High Diversity + Low Poverty', '52', '+0.407', '60%', '27%'],
        ['High Diversity + High Poverty', '73', '+0.808', '59%', '33%'],
    ]
)
doc.add_paragraph(
    'The message is clear: diversity is the switch, not poverty. '
    'Crossing the diversity threshold triples the high-volatility rate (from ~20% to ~60%), '
    'and this holds regardless of whether the county is rich or poor. '
    'Poverty amplifies how much volatile counties swing, but it doesn\'t determine whether they swing.'
)

add_figure(doc, 'deepdive_diversity_poverty_interaction.png',
           'Figure 26: Diversity \u00d7 Poverty quadrant analysis. The left-to-right jump (low to high diversity) '
           'triples volatility rate regardless of poverty level.')

bold_para(doc, 'Education \u00d7 Urbanization \u2014 Volatile at the Extremes:')
add_table(doc,
    ['Quadrant', 'Counties', '% High Vol'],
    [
        ['Low Education + Rural', '89', '42%'],
        ['Low Education + Urban', '36', '19% (least volatile)'],
        ['High Education + Rural', '36', '33%'],
        ['High Education + Urban', '89', '49% (most volatile)'],
    ]
)
doc.add_paragraph(
    'The least volatile counties are low-education urban areas \u2014 stable working-class cities '
    'with long-established voting patterns. The most volatile are high-education urban/suburban '
    'areas \u2014 the fast-growing, college-educated suburbs where political preferences are still forming.'
)

add_figure(doc, 'deepdive_education_urbanization.png',
           'Figure 27: Education \u00d7 Urbanization quadrant analysis.')

bold_para(doc, 'Rent and Volatility \u2014 A U-Shaped Relationship:')
doc.add_paragraph(
    'The rent-volatility relationship is not a straight line. Low-rent counties show mixed volatility. '
    'Mid-rent counties ($750\u2013$1,100) are the most stable \u2014 the comfortable suburban core. '
    'High-rent counties (>$1,100) are the most volatile \u2014 these are gentrifying, fast-changing areas '
    'on the urban fringe.'
)

add_figure(doc, 'deepdive_rent_volatility.png',
           'Figure 28: Rent vs volatility (3-panel). Reveals the U-shaped nonlinear relationship \u2014 '
           'extremes of rent are both associated with higher volatility.')

add_takeaway_box(doc,
    'The interactions tell a clear story: diversity is the switch that determines whether a county is volatile. '
    'Poverty amplifies the magnitude but doesn\'t flip the switch. Educated suburbs are the most volatile; '
    'working-class cities are the most stable. High-rent gentrifying areas are electoral battlegrounds.')

doc.add_heading('13.6 SHAP Case Studies \u2014 Why Specific Counties Swing', level=2)
doc.add_paragraph(
    'SHAP waterfall plots decompose individual county predictions into per-feature contributions. '
    'We examine four archetypal counties to understand distinct volatility mechanisms:'
)
add_table(doc,
    ['County', 'Archetype', '#1 Driver', 'The Story'],
    [
        ['Bucks, PA', 'Suburban swing', 'High rent (+0.85)', 'Large, expensive Philadelphia suburb.\nNear-zero margin. Classic bellwether.'],
        ['Leelanau, MI', 'Demographic misfit', 'Education (+1.50)', 'Highly educated but rural and white.\nVotes Dem against all expectations.'],
        ['Scotland, NC', 'Diverse-poor rural', 'Diversity (+1.25)', 'Highest volatility in dataset.\nDriven by diversity + poverty.'],
        ['Philadelphia, PA', 'Urban counter-example', 'Living alone (\u22120.80)', 'Dense and diverse but STABLE.\nUrban density overrides diversity signal.'],
    ]
)
doc.add_paragraph(
    'Philadelphia is the critical counter-example. It is extremely diverse, high-poverty \u2014 '
    'demographics that "should" predict high volatility. But it is stable. The SHAP waterfall reveals '
    'why: very high rates of solo-living residents and low homeownership create a consistent urban '
    'electorate. Dense urban environments produce voting consistency that overrides the diversity effect.'
)

add_figure(doc, 'deepdive_shap_case_studies.png',
           'Figure 29: SHAP waterfall plots for 4 county archetypes. Each bar shows how one feature '
           'pushes the prediction toward high (right) or low (left) volatility.')

add_takeaway_box(doc,
    'Different counties are volatile for different reasons: suburbs swing because of high rent '
    'and demographic transition, rural misfits swing because of education-driven cultural shift, '
    'and diverse-poor rural counties swing because of unresolved racial/economic tension. '
    'Dense urban areas are the exception \u2014 diversity does not cause volatility in cities.')

doc.add_heading('13.7 SHAP-Based County Archetypes \u2014 Data-Driven County Types', level=2)

add_method_explainer(doc, 'SHAP-Based Clustering',
    'Instead of grouping counties by their raw demographics (which would just separate urban from rural), '
    'we group them by why the model predicts them as volatile or stable. Two counties might look different '
    'demographically but be volatile for the same reason. This approach reveals functionally distinct '
    'county types based on their volatility mechanism.')

add_table(doc,
    ['Cluster', 'n', '% High Vol', 'Key Driver', 'States', 'Archetype Name'],
    [
        ['0', '35', '100%', 'Diversity (1.54)', 'NC (33/35)', 'NC Diverse Rural'],
        ['1', '47', '0%', 'Small population (0.82)', 'NC + MI', 'Stable Small Counties'],
        ['2', '33', '91%', 'Immigration (0.66)', 'MI (21/33)', 'MI Working-Class Transition'],
        ['3', '92', '0%', 'Low diversity (1.07)', 'All states', 'Stable Core'],
        ['4', '43', '81%', 'Large population (0.92)', 'PA+NC+MI', 'Suburban Battlegrounds'],
    ]
)

doc.add_paragraph(
    'Three clusters show near-perfect separation: Cluster 0 (100% volatile), '
    'Cluster 1 (0% volatile), and Cluster 3 (0% volatile). The model finds clean boundaries '
    'between fundamentally different types of counties.'
)

doc.add_paragraph(
    'The most important insight: Cluster 0 (NC Diverse Rural) is almost entirely North Carolina, '
    'and Cluster 2 (MI Working-Class Transition) is mostly Michigan. These are the state-specific '
    'volatility mechanisms that don\'t transfer \u2014 which is why cross-state prediction fails.'
)

add_figure(doc, 'deepdive2_shap_archetypes.png',
           'Figure 30: SHAP archetypes. Top-left: PCA projection of SHAP space (each dot is a county). '
           'Top-right: state composition per cluster. Bottom: radar charts of SHAP feature profiles.')

add_takeaway_box(doc,
    'Five distinct types of counties emerge: NC diverse rural (always volatile), MI working-class '
    'transition (mostly volatile), suburban battlegrounds (mostly volatile), stable small counties '
    '(never volatile), and the stable core (never volatile). These archetypes explain why '
    'cross-state models fail \u2014 each state has its own volatility mechanism.')

doc.add_heading('13.8 Campaign Priority Scoring \u2014 Where to Invest Resources', level=2)
doc.add_paragraph(
    'Translating our analytical findings into a practical tool for campaign resource allocation. '
    'The priority score combines four factors:'
)
doc.add_paragraph('Volatility (30%): How much does this county swing?', style='List Bullet')
doc.add_paragraph('Misfit score (25%): Does it defy demographic expectations?', style='List Bullet')
doc.add_paragraph('Vote volume (25%): How many total votes are cast?', style='List Bullet')
doc.add_paragraph('Margin closeness (20%): How competitive is it?', style='List Bullet')

add_table(doc,
    ['Rank', 'County', 'State', 'Priority', 'Total Votes', '2024 Margin', 'Volatility'],
    [
        ['1', 'Bucks', 'PA', '0.650', '802,056', '\u22120.1%', 'Highly Volatile'],
        ['2', 'Lackawanna', 'PA', '0.574', '233,180', '+2.8%', 'Highly Volatile'],
        ['3', 'Cabarrus', 'NC', '0.570', '120,202', '\u22127.7%', 'Highly Volatile'],
        ['4', 'Leelanau', 'MI', '0.566', '17,685', '+7.8%', 'Highly Volatile'],
        ['5', 'Scotland', 'NC', '0.561', '14,626', '\u22126.9%', 'Highly Volatile'],
        ['6', 'Marquette', 'MI', '0.547', '39,009', '+8.7%', 'Highly Volatile'],
        ['7', 'Genesee', 'MI', '0.514', '223,268', '+4.2%', 'Volatile'],
        ['8', 'Monroe', 'PA', '0.492', '171,050', '\u22120.8%', 'Highly Volatile'],
        ['9', 'Isabella', 'MI', '0.492', '30,835', '\u22127.5%', 'Volatile'],
        ['10', 'Grand Traverse', 'MI', '0.484', '62,772', '\u22121.7%', 'Highly Volatile'],
    ]
)

bold_para(doc, 'Bucks County, PA is the #1 target nationally:',
          ' 802,000 votes, a margin of just 0.1%, and highly volatile. '
          'This single county could decide Pennsylvania, and Pennsylvania could decide the presidency.')

add_figure(doc, 'deepdive_campaign_priority_scatter.png',
           'Figure 31: Priority scatter. X-axis = volatility, Y-axis = misfit score. '
           'Bubble size = vote volume. Top-right = highest priority targets.')

add_takeaway_box(doc,
    'Bucks County PA is the single most important county in the three-state dataset: '
    'massive vote volume, razor-thin margin, and high volatility. PA\'s suburban collar '
    '(Bucks, Chester, Montgomery, Lehigh, Northampton) is the decisive battleground. '
    'MI targets split between small misfit-driven indicators (Leelanau, Marquette) '
    'and large population centers (Macomb, Oakland, Kent).')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 14. GRAND SYNTHESIS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('14. Key Findings & Grand Synthesis', level=1)

doc.add_heading('The Volatility Formula \u2014 What Makes a County Swing', level=2)
doc.add_paragraph(
    'After four notebooks and thousands of model runs, a clear picture emerges. '
    'Electoral volatility is not random \u2014 it follows a specific pattern driven by '
    'the interaction of three demographic factors:'
)

bold_para(doc, '1. Racial/ethnic diversity is the necessary condition.',
          ' The diversity index (race_entropy_norm) is the #1 predictor across all models, '
          'all importance measures, and all states. Low-diversity counties are almost never volatile. '
          'The threshold is 0.55 \u2014 once no single racial group exceeds ~60% of the population, '
          'the odds of high volatility jump 7\u00d7. This is the "switch" that must be flipped.')

bold_para(doc, '2. Economic transition is the amplifier.',
          ' Rising rents, persistent poverty, and shifting education levels amplify volatility '
          'in diverse counties. Counties that are both diverse and economically transitioning '
          'show the highest volatility. But economic factors alone (without diversity) do not produce volatility.')

bold_para(doc, '3. State-level context is the moderator.',
          ' The same demographics produce different political outcomes in different states. '
          'PA volatility is diversity-driven, MI is education-driven, NC is poverty-driven. '
          'National-level models fail because they cannot capture these state-specific political cultures.')

doc.add_heading('What Does NOT Predict Volatility', level=2)
doc.add_paragraph('Poverty alone: Low-diversity poor counties are among the most stable in our dataset', style='List Bullet')
doc.add_paragraph('Education alone: Only predicts volatility when combined with urbanization', style='List Bullet')
doc.add_paragraph('Demographics from other states: Models trained on one state fail completely on another', style='List Bullet')
doc.add_paragraph('Swing direction: The misfit score predicts instability but not which way a county leans', style='List Bullet')

doc.add_heading('The Grand Synthesis Figure', level=2)
doc.add_paragraph(
    'This 6-panel figure summarizes the entire analysis in one image \u2014 from the universal '
    'diversity switch (A) to state-specific mechanisms (B), model performance (C), '
    'trajectory patterns (D), data-driven archetypes (E), and actionable campaign targets (F):'
)

add_figure(doc, 'deepdive2_grand_synthesis.png',
           'Figure 32: Grand Synthesis. A: Diversity switch operates in all 3 states. '
           'B: Different features matter in each state. C: Binary >> 5-class >> ensembles. '
           'D: Blue Bounceback is most common; Steady Red Drift is most volatile. '
           'E: 5 county archetypes from SHAP clustering. F: Campaign priority targets.',
           width=Inches(7.0))

doc.add_heading('Geographic Visualization', level=2)

add_figure(doc, 'choropleth_volatility_class.png',
           'Figure 33: Map of predicted volatility class across PA, MI, and NC. '
           'Red/orange counties are predicted as Volatile or Highly Volatile.')

add_figure(doc, 'choropleth_misfit_score.png',
           'Figure 34: Map of demographic misfit scores. '
           'Darker colors = counties whose voting most defies their demographics.')

doc.add_heading('Complete Metrics At-a-Glance', level=2)
add_table(doc,
    ['Finding', 'Metric', 'Value'],
    [
        ['Best binary model', 'RF F1', '0.718'],
        ['Best 5-class model', 'XGBoost F1-macro', '0.387'],
        ['Bootstrap CI (RF binary)', 'F1 [95% CI]', '0.818 [0.764, 0.872]'],
        ['Performance ceiling', 'Ensemble F1', '0.691 (no improvement)'],
        ['Diversity threshold', 'race_entropy_norm', '0.552 [0.397, 0.642]'],
        ['Diversity odds ratio', 'Above/below threshold', '7.15\u00d7'],
        ['Most necessary group', 'Race/Ethnicity F1 drop', '+0.092'],
        ['Most sufficient group', 'Race/Ethnicity F1 alone', '0.619 (86% of full)'],
        ['Top temporal predictor', '\u0394 pct_hispanic \u03c1', '+0.220'],
        ['Most volatile trajectory', 'Steady Red Drift', '50% high vol'],
        ['Most common trajectory', 'Blue Bounceback', '117/250 (47%)'],
        ['#1 priority county', 'Bucks County PA', '802k votes, 0.1% margin'],
        ['SHAP archetypes', 'k=5 clusters', '100%/0%/91%/0%/81%'],
        ['Cross-state transfer', 'LOSO F1', '0.195 (= random)'],
        ['H1 (income \u2194 volatility)', 'SHAP r', '\u22120.602'],
        ['H2 (entropy > race)', 'Permutation #1', 'Confirmed'],
        ['Misfit-volatility', 'Spearman \u03c1', '0.408 (p < 0.001)'],
        ['Partisan prediction', 'LR accuracy', '89.6%'],
    ]
)

add_takeaway_box(doc,
    'Racial diversity is the master switch for electoral volatility. Economic transition amplifies it. '
    'State context moderates it. National models fail. The best model (RF binary) achieves F1 = 0.718, '
    'and this performance is robust across 200 bootstrap iterations.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 15. PRACTICAL PLAYBOOK
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('15. Practical Playbook \u2014 Democratic Campaign Strategy', level=1)

doc.add_paragraph(
    'This section translates our analytical findings into actionable guidance '
    'for Democratic campaign resource allocation in Pennsylvania, Michigan, and North Carolina. '
    'The recommendations are grounded in data, not intuition.'
)

doc.add_heading('15.1 The Core Strategic Insight', level=2)
doc.add_paragraph(
    'Our analysis reveals that electoral volatility is predictable from demographics. '
    'This means campaigns don\'t need to wait for election results or polling data '
    'to identify swing targets \u2014 census data alone can identify high-leverage counties '
    'years before an election. The key insight: look for racially diverse, economically '
    'transitioning communities \u2014 especially suburbs with rising housing costs.'
)

doc.add_heading('15.2 How to Identify High-Priority Counties', level=2)
doc.add_paragraph(
    'Our priority scoring formula weights four factors. Here is how a campaign analyst '
    'can apply it using publicly available census data:'
)

bold_para(doc, 'Step 1 \u2014 Screen by Diversity: ',
          'Check if the county\'s racial composition has no single group above ~60%. '
          'If one group dominates, the county is almost certainly stable and not worth contesting aggressively. '
          'Counties above this diversity threshold are 7\u00d7 more likely to swing.')

bold_para(doc, 'Step 2 \u2014 Check for Economic Transition: ',
          'Among diverse counties, prioritize those where rents are rising, '
          'education levels are shifting, or there is significant immigration. '
          'These "transitioning" counties have amplified volatility.')

bold_para(doc, 'Step 3 \u2014 Weight by Vote Volume and Margin: ',
          'A volatile county with 800,000 votes and a 0.1% margin (Bucks, PA) '
          'is worth more than a volatile county with 15,000 votes and a 7% margin (Scotland, NC). '
          'Both are volatile, but Bucks delivers 50\u00d7 more votes in a competitive contest.')

bold_para(doc, 'Step 4 \u2014 Tailor by State: ',
          'The same demographics mean different things in different states. '
          'Use state-specific models, not national ones.')

doc.add_heading('15.3 State-by-State Strategy', level=2)

bold_para(doc, 'PENNSYLVANIA \u2014 The Philadelphia Suburban Collar')
doc.add_paragraph(
    'PA\'s volatility is driven by racial diversity and housing cost in the suburban belt. '
    'The decisive battleground is the Philadelphia collar: Bucks (0.1% margin, 802k votes), '
    'Chester, Montgomery, plus the Lehigh Valley (Northampton, Lehigh). '
    'These are high-rent, diversifying suburbs where the electorate is genuinely in flux. '
    'Monroe County (171k votes, 0.8% margin) is an emerging target as NYC-area migration '
    'transforms the Poconos.'
)
doc.add_paragraph('Primary targets: Bucks, Monroe, Northampton, Chester, Lehigh', style='List Bullet')
doc.add_paragraph('Resource strategy: Ground-game investment in diverse suburbs; persuasion-focused messaging targeting economically anxious suburban voters', style='List Bullet')
doc.add_paragraph('Warning: Philadelphia itself is stable despite being diverse \u2014 don\'t conflate suburban and urban dynamics', style='List Bullet')

bold_para(doc, 'MICHIGAN \u2014 Two-Track Targeting')
doc.add_paragraph(
    'MI has a split strategy. Track 1: Small "misfit" counties (Leelanau, Marquette, Isabella, '
    'Grand Traverse) are university/resort towns that vote Democratic against their rural '
    'demographics \u2014 they are volatile and serve as early-warning indicators of broader trends. '
    'Track 2: Large population centers (Macomb, Oakland, Kent/Grand Rapids, Genesee/Flint) '
    'deliver the actual votes. MI volatility is education-driven, not diversity-driven \u2014 '
    'target messaging should emphasize education and economic opportunity.'
)
doc.add_paragraph('Primary targets (volume): Macomb, Oakland, Kent, Genesee', style='List Bullet')
doc.add_paragraph('Primary targets (signal): Leelanau, Marquette, Grand Traverse, Isabella', style='List Bullet')
doc.add_paragraph('Resource strategy: Monitor misfit counties as bellwethers; invest heavily in the population centers', style='List Bullet')
doc.add_paragraph('Key differentiator: Education, not diversity, is the swing driver in MI', style='List Bullet')

bold_para(doc, 'NORTH CAROLINA \u2014 The Southern Realignment')
doc.add_paragraph(
    'NC is undergoing a deeper realignment than PA or MI. Nearly half of NC counties (49%) '
    'show Steady Red Drift \u2014 a sustained move toward Republicans across both election cycles. '
    'Cabarrus County (120k votes, Charlotte suburb) is the standout target: highly volatile, '
    'growing, and diversifying. The small rural misfits (Scotland, Nash, Lenoir) are '
    'diagnostically interesting but individually less decisive. NC\'s volatility is poverty-driven \u2014 '
    'messaging should address economic conditions in diverse, struggling communities.'
)
doc.add_paragraph('Primary targets: Cabarrus (Charlotte suburb), Wake adjacent counties', style='List Bullet')
doc.add_paragraph('Diagnostic targets: Scotland, Nash, Lenoir (small but indicative of broader patterns)', style='List Bullet')
doc.add_paragraph('Resource strategy: Focus on Charlotte-area suburbs; acknowledge the broader red drift but target the exceptions', style='List Bullet')
doc.add_paragraph('Key differentiator: Poverty, not education or diversity alone, drives NC volatility', style='List Bullet')

doc.add_heading('15.4 What NOT to Do', level=2)
doc.add_paragraph('Don\'t apply a one-size-fits-all model: Demographics mean different things in different states. What predicts volatility in PA fails in MI.', style='List Bullet')
doc.add_paragraph('Don\'t target low-diversity counties: They are almost never volatile, regardless of other demographics. Resources spent there are wasted.', style='List Bullet')
doc.add_paragraph('Don\'t conflate urban and suburban dynamics: Dense cities (Philadelphia, Detroit) are stable even when diverse. Volatility lives in the suburbs and exurbs.', style='List Bullet')
doc.add_paragraph('Don\'t assume the 2020 surge is permanent: 47% of counties showed "Blue Bounceback" \u2014 the 2020 gains were temporary. Plan for reversion, not continuation.', style='List Bullet')
doc.add_paragraph('Don\'t ignore "Steady Red Drift" counties: These are the most volatile (50% high vol) and represent an ongoing realignment, especially in NC.', style='List Bullet')

doc.add_heading('15.5 The "Money List" \u2014 10 Counties That Could Decide 2028', level=2)
doc.add_paragraph(
    'Based on our composite priority scoring (volatility + misfit + votes + margin), '
    'these 10 counties offer the highest marginal utility for Democratic investment:'
)
add_table(doc,
    ['Rank', 'County', 'State', 'Why It Matters'],
    [
        ['1', 'Bucks', 'PA', '802k votes, 0.1% margin, diverse suburb \u2014 could decide PA'],
        ['2', 'Lackawanna', 'PA', '233k votes, Scranton area, blue misfit trending volatile'],
        ['3', 'Cabarrus', 'NC', '120k votes, Charlotte suburb, fast-growing and diversifying'],
        ['4', 'Leelanau', 'MI', 'Small but highest misfit score \u2014 bellwether for MI trends'],
        ['5', 'Scotland', 'NC', 'Highest volatility score, diverse-poor rural archetype'],
        ['6', 'Marquette', 'MI', 'University town, votes Dem against all demographics'],
        ['7', 'Genesee', 'MI', '223k votes (Flint area), working-class transition zone'],
        ['8', 'Monroe', 'PA', '171k votes, Poconos, NYC migration transforming politics'],
        ['9', 'Isabella', 'MI', 'University county (CMU), young and volatile'],
        ['10', 'Grand Traverse', 'MI', '63k votes, resort/wine country, 1.7% margin, highly volatile'],
    ]
)

add_takeaway_box(doc,
    'For maximum impact: invest in PA\'s suburban collar (Bucks, Monroe, Lackawanna), '
    'MI\'s population centers (Genesee, Oakland, Macomb) while monitoring misfit bellwethers '
    '(Leelanau, Marquette), and NC\'s Charlotte-area suburbs (Cabarrus). '
    'Use diversity screening as the first filter: if a county isn\'t racially mixed, '
    'it probably isn\'t going to swing.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 16. ANSWERING ALL RESEARCH QUESTIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('16. Answering All Research Questions', level=1)

doc.add_paragraph(
    'This section maps our findings directly to each research question from the project proposal.'
)

doc.add_heading('Results Application Questions', level=2)

bold_para(doc, 'Q: Which counties illustrate a combination of electoral elasticity and impact (vote volume), '
               'maximizing the marginal utility of Democratic resource allocation?')
doc.add_paragraph(
    'Our campaign priority scoring (Section 13.8) directly answers this. Bucks County, PA is #1: '
    '802,000 votes, 0.1% margin, Highly Volatile. The full top 10 combines volatility, misfit score, '
    'vote volume, and margin closeness into a single priority ranking. '
    'Among high-volume counties (>100k votes), Bucks, Lackawanna, Cabarrus, Genesee, and Monroe '
    'offer the best return on campaign investment. See Section 15.5 for the complete "Money List."'
)

bold_para(doc, 'Q: Which swing-state counties had the highest electorate volatility between 2016, 2020, and 2024?')
doc.add_paragraph(
    'By our vol_z_abs_sum measure, the most volatile counties include Scotland (NC), '
    'Leelanau (MI), Marquette (MI), Bucks (PA), Cabarrus (NC), and Lackawanna (PA). '
    'The trajectory analysis (Section 13.2) shows that "Steady Red Drift" counties are the most '
    'volatile as a group (50% high volatility, mean vol = +0.39), while "Blue Bounceback" '
    'counties (the 2020 Biden surge that reversed in 2024) are the most common pattern (47% of counties). '
    'In terms of vote volume, Bucks PA (802k), Genesee MI (223k), and Lackawanna PA (233k) '
    'represent the highest-volume volatile counties.'
)

bold_para(doc, 'Q: What specific demographic and economic features distinguish "High-Volatility" '
               'county clusters from "Stable" clusters?')
doc.add_paragraph(
    'The feature ablation analysis (Section 10.4) and SHAP analysis (Section 5) provide a definitive answer. '
    'High-volatility counties are distinguished by: (1) Higher racial diversity index (the #1 predictor, '
    'odds ratio 7.15\u00d7 above threshold), (2) Higher housing costs (rent is the only universal predictor '
    'across all 3 states), (3) Higher poverty rates (especially in NC), (4) Higher education levels '
    '(especially in MI), and (5) Lower homeownership and younger median age. '
    'Stable counties are the opposite: racially homogeneous, moderate housing costs, higher homeownership, '
    'and older populations with established voting patterns. '
    'The SHAP archetype analysis (Section 13.7) identifies 5 distinct cluster types, '
    'with near-perfect separation between volatile and stable groups.'
)

doc.add_heading('Hypotheses / Predictions', level=2)

bold_para(doc, 'H1: Electoral volatility is inversely correlated with wealth.')
doc.add_paragraph(
    'CONFIRMED. The SHAP analysis of income features shows r = \u22120.602 (p < 0.0001) between '
    'household income and high-volatility SHAP values. Poverty rate is a top-5 predictor '
    'across all models. However, the relationship is nuanced: income alone is insufficient '
    '(economic features alone achieve only F1 = 0.570). Income interacts with diversity \u2014 '
    'wealthy diverse suburbs are volatile, while wealthy homogeneous exurbs are not. '
    'See Section 12 for full evidence.'
)

bold_para(doc, 'H2: A composite Racial Diversity Index will demonstrate higher feature importance '
               'than any single demographic.')
doc.add_paragraph(
    'CONFIRMED across 6 independent evidence streams: (1) #1 in permutation importance, '
    '(2) higher SHAP than pct_black (0.397 vs 0.347), (3) largest PDP range (0.14), '
    '(4) ablation shows race group alone achieves 86% of full performance, '
    '(5) appears in 6/10 top SHAP interactions, (6) threshold analysis shows 7.15\u00d7 odds ratio. '
    'The diversity index captures information about racial mixing that no single-race variable provides. '
    'See Section 12 for complete evidence table.'
)

bold_para(doc, 'H3: High-volatility counties function as macro-trend amplifiers.')
doc.add_paragraph(
    'PARTIALLY CONFIRMED, with important caveats. The trajectory analysis (Section 13.2) shows that '
    'the dominant trajectory is "Blue Bounceback" (47%) \u2014 counties that swung Democratic in 2020 '
    'and reverted in 2024, amplifying both the 2020 blue wave and the 2024 correction. '
    '"Steady Red Drift" counties (38%) amplify the Republican trend with greater magnitude '
    '(mean vol = +0.39, 50% high volatility). However, we cannot confirm that volatile counties '
    'always follow the statewide trend \u2014 the cross-state analysis shows volatility mechanisms '
    'are state-specific, and the misfit analysis shows some counties swing against expectations. '
    'The amplifier hypothesis holds for the majority of counties but is not universal.'
)

doc.add_heading('Further Contextual Analysis Questions', level=2)

bold_para(doc, 'Q: To what extent do demographic predictors generalize across state lines?')
doc.add_paragraph(
    'They DO NOT generalize. This is one of our most striking findings (Section 11). '
    'Leave-One-State-Out cross-validation produces F1 = 0.195, indistinguishable from random '
    'guessing (0.200). The state-specific analysis reveals why: PA volatility is diversity-driven, '
    'MI is education-driven, NC is poverty-driven. Only housing cost (rent) predicts volatility '
    'in all three states. The SHAP archetype analysis confirms this: Cluster 0 (NC Diverse Rural, '
    '33/35 NC counties) and Cluster 2 (MI Working-Class Transition, 21/33 MI counties) '
    'represent state-specific mechanisms that cannot transfer.'
)

bold_para(doc, 'Q: Does electoral volatility exhibit temporal consistency? '
               'Does the "Swing Map" shift between cycles?')
doc.add_paragraph(
    'The temporal analysis (Section 13.1\u201313.2) addresses this directly. The swing map does shift '
    'significantly between cycles. 47% of counties showed "Blue Bounceback" \u2014 they swung blue '
    'in 2020 and red in 2024, meaning the geographic composition of the swing map changed substantially. '
    'Only 14% of counties ("Steady Blue Drift") swung consistently in one direction across both cycles. '
    'Furthermore, the demographic change analysis shows that counties with growing Hispanic populations '
    'and rising rents (\u03c1 = +0.22, +0.21) are becoming more volatile over time, '
    'suggesting the swing map is dynamically evolving with demographic change, not static.'
)

bold_para(doc, 'Q: What is the statistical correlation between voter turnout variance and vote share volatility?')
doc.add_paragraph(
    'Our analysis focused on vote share volatility (margin changes) rather than turnout variance, '
    'as the research design prioritized demographic predictors of partisan swing. '
    'However, we can provide indirect evidence: (1) The misfit analysis uses total_votes as a '
    'component of the priority score, and there is no strong correlation between total vote volume '
    'and volatility in our data \u2014 both large and small counties can be volatile. '
    '(2) The log_total_population feature ranks moderate in importance but works primarily through '
    'interactions (PDP is mostly flat), suggesting that turnout effects are context-dependent. '
    'A direct turnout-variance analysis would require year-over-year turnout data, which could be '
    'explored in future work as an extension of this framework.'
)

add_takeaway_box(doc,
    'All primary research questions are answered. Both hypotheses confirmed. '
    'Cross-state generalization fails (politics is local). The swing map shifts between cycles. '
    'High-volatility counties generally amplify macro trends but with important exceptions. '
    'The diversity index is the master predictor, operating as a threshold switch at 0.55.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 17. APPENDIX
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('17. Appendix: Notebooks, Data Files & Figures', level=1)

doc.add_heading('Notebooks', level=2)
add_table(doc,
    ['Notebook', 'Cells', 'Description'],
    [
        ['classification_antoni.ipynb', '86', 'Core: RF, XGBoost, SVM, Misfit Detector, LOSO, maps'],
        ['classification_antoni_deepdive.ipynb', '48', 'Deep-dive I: misfits, state models, SHAP interactions,\nbinary models, campaign scoring, nonlinear analysis'],
        ['classification_antoni_deepdive_2.ipynb', '30', 'Deep-dive II: temporal, trajectories, PDPs,\nbootstrap, ablation, ensembles, threshold, archetypes'],
        ['classification_antoni_deepdive_3.ipynb', '5', 'Grand synthesis figure (standalone)'],
    ]
)

doc.add_heading('Data Files', level=2)
add_table(doc,
    ['File', 'Description'],
    [
        ['classification_dataset_250.csv', '250 counties, 28 scaled features + targets'],
        ['classification_predictions.csv', 'All model predictions + misfit scores'],
        ['rf_predictions.csv', 'RF predictions for clustering comparison'],
        ['campaign_priority_rankings.csv', '250 counties with priority scores'],
    ]
)

doc.add_heading('Figure Index (34 figures)', level=2)
figures_list = [
    ('1', 'RF confusion matrix (5-class)'),
    ('2', 'RF Gini feature importance'),
    ('3', 'RF permutation importance'),
    ('4', 'XGBoost confusion matrix (5-class)'),
    ('5', 'SHAP feature importance summary'),
    ('6', 'SHAP dependence plots (top 5)'),
    ('7', 'Top 6 SHAP interactions'),
    ('8', 'SHAP interaction heatmap'),
    ('9', 'SVM confusion matrix'),
    ('10', 'P(Dem) vs actual margin scatter'),
    ('11', 'Cross-analysis correlations'),
    ('12', 'Misfit score by volatility class'),
    ('13', 'Binary ROC curves (5 models)'),
    ('14', 'Binary confusion matrices'),
    ('15', 'SHAP summary (binary)'),
    ('16', 'Feature ablation + permutation CIs'),
    ('17', 'Final model scoreboard'),
    ('18', 'Per-class F1 heatmap'),
    ('19', 'LOSO confusion matrices'),
    ('20', 'Per-state feature importance'),
    ('21', 'State-specific ROC curves'),
    ('22', 'Temporal demographic shifts'),
    ('23', 'Margin trajectories (spaghetti)'),
    ('24', 'Diversity threshold'),
    ('25', 'Partial dependence plots'),
    ('26', 'Diversity \u00d7 Poverty interaction'),
    ('27', 'Education \u00d7 Urbanization'),
    ('28', 'Rent vs volatility (nonlinear)'),
    ('29', 'SHAP case studies (4 counties)'),
    ('30', 'SHAP-based archetypes'),
    ('31', 'Campaign priority scatter'),
    ('32', 'Grand synthesis (6-panel)'),
    ('33', 'Choropleth: volatility class'),
    ('34', 'Choropleth: misfit score'),
]
add_table(doc,
    ['Figure', 'Description'],
    [[f[0], f[1]] for f in figures_list]
)

# ── Save ──
doc.save(OUT)
print(f'Document saved to: {OUT}')
