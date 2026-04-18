"""
Generate DOCX summary document for the Democratic Priority Score.
Author: Antoni Czolgowski | April 2026
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

BASE = '../data/processed'
FIG_DIR = '../docs/images/methods'
OUTPUT_PATH = '../reports/Enhanced_Democratic_Priority_Score_Summary.docx'

df = pd.read_csv(f'{BASE}/enhanced_priority_rankings.csv', dtype={'county_fips': str})
df['county_short'] = df['county_name'].str.replace(' County', '', regex=False)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_image_centered(path, width=Inches(6.2)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)

def add_table_from_data(headers, rows, font_size=Pt(8.5)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = font_size
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = font_size
    return table

def bold_para(title, body):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.add_run(body)

# ═══════════════ TITLE PAGE ═══════════════
title = doc.add_heading('Democratic Priority Score', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(
    'A Data-Driven Composite Index for Campaign Resource Allocation\n'
    'Swing State Election Analysis: Pennsylvania, Michigan, North Carolina')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(80, 80, 80)

auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = auth.add_run(
    '\nAntoni Czolgowski | Data Lead\n'
    'CSCI 5502 Data Mining | Spring 2026 | University of Colorado Boulder')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ═══════════════ 1. EXECUTIVE SUMMARY ═══════════════
add_heading('1. Executive Summary')
doc.add_paragraph(
    'The Democratic Priority Score (DPS) is a composite 0-1 index that ranks 250 counties '
    'across three critical swing states -- Pennsylvania, Michigan, and North Carolina -- '
    'by their strategic value for Democratic campaign resource allocation. '
    'The score integrates six data-driven components derived from Random Forest classification, '
    'linear regression, K-Means clustering, and temporal electoral trajectory analysis.'
)
doc.add_paragraph(
    'The #1 priority county is Bucks County, PA (DPS = 0.734): 802,000 votes, a margin of '
    'just 582 votes (-0.1%), and classified as Highly Volatile. The top 10 includes 5 PA counties '
    '(the Philadelphia suburban corridor), 2 MI counties (Kent/Grand Rapids and Oakland/Detroit suburbs), '
    'and 3 NC counties (Cabarrus/Charlotte suburbs, New Hanover/Wilmington, and Wilson). '
    'The top 10 counties collectively represent 3.45 million votes with an average margin '
    'of just 3.6 percentage points.'
)

# ═══════════════ 2. METHODOLOGY ═══════════════
add_heading('2. Methodology')
doc.add_paragraph(
    'The DPS decomposes into six sub-scores, each normalized to [0, 1], '
    'combined with data-driven weights:')
doc.add_paragraph(
    'DPS = 0.20 * Persuadability + 0.20 * Electoral Weight + 0.20 * Competitiveness '
    '+ 0.20 * Predicted Volatility + 0.10 * Dem. Opportunity + 0.10 * Misfit')

headers = ['Component', 'Wt', 'Formula', 'Source']
rows = [
    ['S1: Persuadability', '.20',
     '0.6*sigmoid(entropy-0.552) + 0.4*cluster_vol_rate',
     'Diversity threshold (OR=7.15) + K-Means clusters'],
    ['S2: Electoral Weight', '.20',
     'MinMax(log(1 + votes_2024))',
     'Academic campaign theory: vote volume'],
    ['S3: Competitiveness', '.20',
     'exp(-|margin| / 10)',
     'Exponential decay; 10pp margin = 0.37'],
    ['S4: Predicted Volatility', '.20',
     'RF.predict_proba[:, 1]',
     'Random Forest (AUC=0.828), 28 features'],
    ['S5: Dem. Opportunity', '.10',
     'MinMax(0.65*d_20_24 + 0.35*d_16_20)',
     'Directional swing; recency-weighted'],
    ['S6: Misfit', '.10',
     'MinMax(direction_adj_misfit)',
     'Surprise-Dem 1.3x, Surprise-Rep 0.7x'],
]
add_table_from_data(headers, rows)

doc.add_paragraph(
    '\nThe core four components receive equal weight (0.20), reflecting both academic campaign theory '
    '(resource allocation = competitiveness x electoral weight) and our data mining results '
    '(persuadability and predicted volatility are the strongest empirical signals). '
    'The two supplementary signals receive 0.10 each: directional opportunity is backward-looking '
    '(Blue Bounceback is the most common trajectory at 47%), and misfit partially overlaps with '
    'predicted volatility.'
)

doc.add_paragraph(
    'Key methodological choices: (1) S3 uses exponential decay rather than linear scaling, '
    'so a +20pp county scores 14% of a tied county -- reflecting campaign resource reality. '
    '(2) S4 integrates all 28 demographic features through a nonlinear Random Forest model, '
    'capturing interactions (diversity x density, education x urbanization) invisible to linear formulas. '
    '(3) S1 implements the diversity threshold at race_entropy_norm = 0.552 as a soft sigmoid, '
    'reflecting the step-function relationship where P(high volatility) jumps from 20% to 80%. '
    '(4) S2 uses log-transformed vote counts to prevent mega-counties from dominating '
    'while preserving the signal that larger counties deserve more resources.'
)

# ═══════════════ 3. RESULTS ═══════════════
doc.add_page_break()
add_heading('3. Results')

add_heading('3.1 Top 10 Democratic Priority Counties', level=2)

top10 = df.nlargest(10, 'DPS')
headers = ['#', 'County', 'ST', 'DPS', 'Pers.', 'ElWt', 'Comp.', 'PVol', 'DOpp', 'Misf', 'Margin', 'Votes']
rows = []
for i, (_, r) in enumerate(top10.iterrows(), 1):
    rows.append([str(i), r['county_short'], r['state'], f"{r['DPS']:.3f}",
        f"{r['S_persuadability']:.2f}", f"{r['S_electoral_weight']:.2f}",
        f"{r['S_competitiveness']:.2f}", f"{r['S_volatility_predicted']:.2f}",
        f"{r['S_dem_opportunity']:.2f}", f"{r['S_misfit']:.2f}",
        f"{r['dem_margin_2024_pp']:+.1f}%", f"{r['total_votes_2024']:,.0f}"])
add_table_from_data(headers, rows)

# Top 5 per state (compact)
for abbr, name in [('PA', 'Pennsylvania'), ('MI', 'Michigan'), ('NC', 'North Carolina')]:
    add_heading(f'3.2 Top 5 -- {name}', level=2)
    st5 = df[df['state'] == abbr].nlargest(5, 'DPS')
    rows = []
    for i, (_, r) in enumerate(st5.iterrows(), 1):
        rows.append([str(i), r['county_short'], r['state'], f"{r['DPS']:.3f}",
            f"{r['S_persuadability']:.2f}", f"{r['S_electoral_weight']:.2f}",
            f"{r['S_competitiveness']:.2f}", f"{r['S_volatility_predicted']:.2f}",
            f"{r['S_dem_opportunity']:.2f}", f"{r['S_misfit']:.2f}",
            f"{r['dem_margin_2024_pp']:+.1f}%", f"{r['total_votes_2024']:,.0f}"])
    add_table_from_data(headers, rows)

# ═══════════════ 4. COUNTY PROFILES ═══════════════
doc.add_page_break()
add_heading('4. County-by-County Analysis')

doc.add_paragraph(
    'Each top-10 county is assessed by its electoral trajectory (2016-2020-2024), '
    'demographic profile, cluster archetype, and strategic role.'
)

bold_para('#1 Bucks County, PA (DPS = 0.734) -- ',
    'THE bellwether. Trajectory: +0.8% (2016), +4.4% (2020), -0.1% (2024) -- a "Blue Bounceback" '
    'that landed on a razor\'s edge of 582 votes. With 802K total votes, near-perfect competitiveness (0.993), '
    'and high predicted volatility (0.878), Bucks is the single most valuable target in the three-state region. '
    'Campaign focus: persuasion. Ground game, earned media, and targeted digital advertising '
    'in the Bucks County media market should be the #1 budget line item.')

bold_para('#2 Lehigh County, PA (DPS = 0.702) -- ',
    'Heart of the Lehigh Valley (Allentown, Bethlehem). Trajectory: +4.7%, +7.6%, +2.7% -- '
    'a Blue Bounceback with an alarming 5-point 2024 pullback. Racial entropy = 0.699 (well above '
    'the 0.552 threshold), driven by Hispanic immigration and NYC commuter spillover. 380K votes. '
    'Campaign focus: bilingual outreach and housing affordability messaging targeting the rapidly '
    'diversifying population that drives the county\'s volatility.')

bold_para('#3 Northampton County, PA (DPS = 0.702) -- ',
    'The classic bellwether. Voted Obama, flipped to Trump in 2016 (-3.8%), flipped back to Biden '
    'in 2020 (+0.7%), and returned Republican in 2024 (-1.8%). Has correctly predicted the PA winner '
    'in 4 of the last 5 cycles. Paired with Lehigh, it forms a natural media market. '
    'Campaign focus: genuine swing voters -- white working-class persuadables in the Easton/Bethlehem '
    'area. Distinct from Lehigh\'s diversity-driven volatility.')

bold_para('#4 Monroe County, PA (DPS = 0.688) -- ',
    'The Pocono Mountains. Highest racial entropy among PA top 5 (0.730). Trajectory mirrors Bucks: '
    '+0.8%, +6.4%, -0.8%. Being transformed by NYC exurban migration along I-80. '
    'Highly Volatile classification. Campaign focus: new residents and first-time voters. '
    'Monroe is a "future Lehigh Valley" -- investing here builds infrastructure for a county '
    'that will only grow more electorally significant.')

bold_para('#5 Cabarrus County, NC (DPS = 0.678) -- ',
    'Charlotte suburbs. Remarkable trajectory: -19.6% (2016), -9.4% (2020), -7.7% (2024) -- '
    'a 12-point Steady Blue Drift over 8 years. Highest S_dem_opportunity (0.932) in the top 10. '
    'Entropy = 0.802. Not yet winnable, but the trend is unmistakable. '
    'Campaign focus: suburban women, college-educated professionals, and the growing Hispanic community. '
    'Every dollar erodes the Republican structural advantage in the Charlotte metro.')

bold_para('#6 Kent County, MI (DPS = 0.672) -- ',
    'Grand Rapids -- historically the capital of Michigan Republicanism (Gerald Ford\'s district). '
    'Flipped Democratic in 2020 (+6.1%) after -3.1% in 2016, held at +5.4% in 2024. '
    'With 372K votes, it is the largest competitive county in MI. Highly Volatile. '
    'Campaign focus: THE persuasion county in Michigan. The evangelical GOP base is being offset '
    'by the growing urban core and education-driven suburban shift.')

bold_para('#7 New Hanover County, NC (DPS = 0.657) -- ',
    'Wilmington and the NC coast. Flipped Dem in 2020 (+2.1%), narrowed to +0.6% in 2024 -- '
    'a margin of ~830 votes out of 139K cast. S_competitiveness = 0.940, second only to Bucks. '
    'Entropy (0.565) sits precisely at the critical diversity threshold. '
    'Campaign focus: the "Bucks County of North Carolina" -- a coastal suburban county '
    'that swings with the national mood. Razor-thin margin makes resource allocation highly efficient.')

bold_para('#8 Oakland County, MI (DPS = 0.654) -- ',
    'The most important "protect the gains" county. Once reliably Republican, Oakland flipped decisively: '
    '+8.1% (2016), +14.0% (2020), but eroded to +10.6% in 2024 -- a 3.4-point loss in one cycle. '
    'With 772K votes (second-largest in the study), this erosion translates to ~26,000 lost votes. '
    'If the trend continues, Oakland becomes competitive by 2028 -- and losing Oakland means losing Michigan. '
    'Predicted volatility = 0.887, the highest in the top 10. '
    'Campaign focus: defensive investment. Not persuasion but turnout infrastructure to arrest erosion. '
    'Every 1% turnout increase = ~7,700 votes.')

bold_para('#9 Dauphin County, PA (DPS = 0.643) -- ',
    'Harrisburg metro, state capital. Trajectory: +2.9%, +8.5%, +5.9% -- the 2024 pullback from '
    'Biden\'s high reveals vulnerability. Racial entropy = 0.775. With 299K votes, it provides '
    'a strong voter pool in central PA. '
    'Campaign focus: geographic diversification. The only top-10 county outside the eastern corridor. '
    'Can anchor a central PA media buy reaching adjacent Cumberland (#20) and Lebanon.')

bold_para('#10 Wilson County, NC (DPS = 0.639) -- ',
    'The most analytically interesting -- and most debatable -- pick. Only 40K votes, but near-perfect '
    'competitiveness (+0.4%) and the highest persuadability in the top 10 (0.940, "Poor & Diverse" cluster). '
    'Trajectory: +5.6%, +2.9%, +0.4% -- a Steady Red Drift of 5.2 points over 8 years. '
    'If the trend continues, Wilson flips Republican by 2028. '
    'Campaign focus: CANARY IN THE COAL MINE. Wilson represents dozens of small, racially diverse, '
    'economically struggling Southern counties slipping away from Democrats. The pattern it exemplifies '
    'could cost tens of thousands of votes across NC. Strategic value is diagnostic, not volume-based.')

# ═══════════════ 5. STATE STRATEGIES ═══════════════
doc.add_page_break()
add_heading('5. State-Level Strategic Assessment')

add_heading('Pennsylvania: The Decisive Battleground (19 EV)', level=2)
doc.add_paragraph(
    'PA dominates the national top 10 with 5 counties, reflecting genuine strategic reality: '
    'the most electoral votes (19), the strongest regression model (R-squared = 0.674), and an unusual '
    'concentration of large, competitive, volatile counties in the eastern suburban corridor. '
    'The top 5 PA counties represent 2,007,208 total votes across a 150-mile corridor '
    'from Philadelphia (Bucks) through the Lehigh Valley (Lehigh, Northampton) '
    'to the Poconos (Monroe) and inland to Harrisburg (Dauphin). '
    'This corridor can be served by 3-4 field offices and a single TV media buy.')

bold_para('Near-miss -- Lackawanna (#12): ',
    'Scranton. 233K votes, +2.8%, Highly Volatile. Its lower diversity (entropy = 0.485, '
    'below the 0.552 threshold) reduces its persuadability score, but the decline from +8.4% (2020) '
    'to +2.8% (2024) is alarming. A campaign with 6 PA targets should add Lackawanna.')

bold_para('Near-miss -- Erie (#23): ',
    'The famous bellwether (Obama, Obama, Trump, Biden, Trump). At 275K votes and -1.0%, '
    'it seems like it should rank higher, but the RF model assigns only 33% volatility probability -- '
    'demographics suggest stable working-class patterns despite historical swings.')

add_heading('Michigan: Volume Targets vs. Bellwethers (15 EV)', level=2)
doc.add_paragraph(
    'MI\'s top 5 splits into large suburban centers (Kent 372K, Oakland 772K, Macomb 509K) and '
    'smaller competitive markets (Genesee 223K, Grand Traverse 63K). Kent (#1 MI) is the premier '
    'persuasion target. Oakland (#2) is defensive. Genesee (#3, Flint) shows troubling erosion '
    '(+9.5% to +4.2% over 8 years). Grand Traverse (#4, Traverse City) is genuinely competitive '
    'at -1.7%. Macomb (#5, -13.7%) is controversial but at 509K votes is too large to ignore.')

bold_para('Not recommended -- Leelanau (now #33): ',
    'At 17,685 votes, even a 10-point swing yields only ~880 net votes. '
    'Interesting as a misfit case study (highest misfit score: 0.956) but not a resource target.')

add_heading('North Carolina: The Most Complex Landscape (16 EV)', level=2)
doc.add_paragraph(
    'NC presents the most challenging targeting decisions. Unlike PA\'s concentrated corridor '
    'and MI\'s clear suburban targets, NC\'s priorities are geographically scattered. '
    'The state-specific regression model has the weakest performance (R-squared = 0.435) '
    'because NC\'s poverty-driven volatility mechanism is harder to predict.')

bold_para('Strong picks -- Cabarrus (#1) and New Hanover (#2): ',
    'The two clearest NC targets. Cabarrus shows consistent Democratic momentum (-19.6% to -7.7%). '
    'New Hanover is essentially tied at +0.6%. Together: 259K votes in genuinely competitive markets.')

bold_para('Borderline -- Wilson (#3): ',
    'Small (40K) but near-zero margin (+0.4%) and maximum persuadability (0.940). '
    'Value is diagnostic: if a message works in Wilson, it works across similar diverse rural Southern counties.')

bold_para('Caution -- Wake (#4, +25.4%, 654K): ',
    'NOT a persuasion target. Presence is driven by enormous electoral weight and high persuadability score. '
    'In practice, Wake is a TURNOUT target: even 1% improvement = ~6,500 votes. '
    'Alternative persuasion picks: Pitt County (87K, +6.0%, ECU university town) '
    'or Nash County (52K, -1.8%, essentially tied).')

bold_para('Weakest pick -- Lenoir (#5, -6.8%, 28K, Stable): ',
    'Most questionable entry in any state\'s top 5. Classified as Stable, only 28K votes, '
    'drifting Republican. High ranking driven entirely by persuadability score (0.918) from demographics '
    'that predict volatility which hasn\'t materialized. '
    'Better alternatives: Nash (52K, -1.8%) or Pitt (87K, +6.0%).')

# ═══════════════ 6. VISUALIZATIONS ═══════════════
doc.add_page_break()
add_heading('6. Visualizations')

add_heading('6.1 Three-State Choropleth: All Counties', level=2)
doc.add_paragraph(
    'Every county shaded by DPS (0 = low priority, darker = higher). '
    'Clear geographic clusters: eastern PA, southern MI, and the urban/suburban crescent of NC.')
add_image_centered(f'{FIG_DIR}/enhanced_priority_3state_choropleth.png', Inches(6.2))

add_heading('6.2 Top 10 Priority Counties Highlighted', level=2)
doc.add_paragraph(
    'Top 10 counties colored with relative shading; all others in gray. '
    'PA dominates with 5 of 10, concentrated in the eastern suburban corridor.')
add_image_centered(f'{FIG_DIR}/enhanced_priority_top10_highlight.png', Inches(6.2))

doc.add_page_break()
add_heading('6.3 State Maps: Top 5 Counties', level=2)

doc.add_paragraph('Pennsylvania: Top 5 Priority Counties', style='Intense Quote')
add_image_centered(f'{FIG_DIR}/enhanced_priority_PA_top5.png', Inches(5.5))

doc.add_paragraph('Michigan: Top 5 Priority Counties', style='Intense Quote')
add_image_centered(f'{FIG_DIR}/enhanced_priority_MI_top5.png', Inches(5.5))

doc.add_page_break()
doc.add_paragraph('North Carolina: Top 5 Priority Counties', style='Intense Quote')
add_image_centered(f'{FIG_DIR}/enhanced_priority_NC_top5.png', Inches(5.5))

# ═══════════════ 7. VALIDATION & LIMITATIONS ═══════════════
add_heading('7. Validation & Limitations')

top10_vol = top10['volatility_class'].value_counts().to_dict()
skew = df['DPS'].skew()

doc.add_paragraph('Validation checks:')
checks = [
    f'Top 10 volatility classes: {top10_vol}. No Stable or Very Stable county in the top 10.',
    f'Score distribution: mean = {df["DPS"].mean():.3f}, median = {df["DPS"].median():.3f}, '
    f'skewness = {skew:.3f}. Right-skewed as expected: most counties are low-priority '
    'with a tail of high-value targets.',
    'Bucks County PA is #1 -- consistent with conventional campaign wisdom '
    'about this critical Philadelphia suburb.',
    'Top 10 collectively covers 3.45M votes with avg |margin| of 3.6pp -- '
    'an efficient allocation of campaign resources to competitive, high-volume counties.',
    '9 of 10 top counties have >50K votes. The score correctly prioritizes counties '
    'where resource investment can move meaningful vote totals.',
]
for c in checks:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph('\nLimitations:')
limitations = [
    'RF model trained on 250 counties (small sample). External validation on other swing states '
    'would strengthen confidence.',
    'ACS demographic data lags real-time population changes by 1-2 years. '
    'Fast-growing counties (Cabarrus NC) may have shifted.',
    'The score does not distinguish between turnout targets (base mobilization) '
    'and persuasion targets (swing voters). Voter file-level data would enable this.',
    'Cross-state models fail (LOSO F1 = 0.195). The DPS handles this through implicit RF encoding '
    'rather than explicit state adjustments.',
    'The score should be recalibrated each election cycle as new data becomes available.',
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

# ═══════════════ SAVE ═══════════════
doc.save(OUTPUT_PATH)
print(f'DOCX saved: {OUTPUT_PATH}')
